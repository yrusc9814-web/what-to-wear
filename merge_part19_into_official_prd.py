from copy import deepcopy
from pathlib import Path

from docx import Document


DOC_DIR = Path(r"D:\产品文档")
BASE_WITH_18 = DOC_DIR / "穿搭日记微信小程序_PRD_v1.0_已加入第18部分.docx"
OFFICIAL = DOC_DIR / "穿搭日记微信小程序_PRD_v1.0.docx"
PART19 = DOC_DIR / "穿搭日记微信小程序_PRD_第19部分_开发实施计划_v1.4_审核稿.docx"
FALLBACK = DOC_DIR / "穿搭日记微信小程序_PRD_v1.0_已加入第18和第19部分.docx"


def element_text(el):
    return "".join(node.text or "" for node in el.iter() if node.tag.endswith("}t"))


def strip_existing_part19(doc):
    body = doc.element.body
    children = list(body)
    start_index = None
    for i, el in enumerate(children):
        if el.tag.endswith("}p") and element_text(el).strip().startswith("19.1 "):
            start_index = i
            break
    if start_index is None:
        return
    for el in children[start_index:]:
        if not el.tag.endswith("}sectPr"):
            body.remove(el)


def append_part19(base_doc, part_doc):
    strip_existing_part19(base_doc)
    base_doc.add_page_break()
    base_body = base_doc.element.body
    insert_before = base_body.sectPr

    include = False
    for el in part_doc.element.body:
        if el.tag.endswith("}sectPr"):
            continue
        if el.tag.endswith("}p") and element_text(el).strip().startswith("19.1 "):
            include = True
        if include:
            base_body.insert(base_body.index(insert_before), deepcopy(el))


def contains_marker(path, marker):
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return marker in text


def main():
    if not BASE_WITH_18.exists():
        raise FileNotFoundError(BASE_WITH_18)
    if not PART19.exists():
        raise FileNotFoundError(PART19)

    base_doc = Document(BASE_WITH_18)
    part_doc = Document(PART19)
    append_part19(base_doc, part_doc)

    try:
        base_doc.save(OFFICIAL)
        target = OFFICIAL
    except PermissionError:
        base_doc.save(FALLBACK)
        target = FALLBACK

    print(target)
    print("HAS_18", contains_marker(target, "18.1"))
    print("HAS_19", contains_marker(target, "19.1"))


if __name__ == "__main__":
    main()
