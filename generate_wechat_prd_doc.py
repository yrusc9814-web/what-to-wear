from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from pathlib import Path


OUTPUT_DIR = Path(r"D:\产品文档")
OUTPUT_FILE = OUTPUT_DIR / "穿搭日记微信小程序_PRD_v1.0.docx"


def set_font(run, name="Microsoft YaHei", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph(document, text="", bold=False, size=11, align=None):
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def add_heading(document, text, level=1):
    return add_paragraph(document, text, bold=True, size=16 if level == 1 else 13)


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run)


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)

    add_paragraph(doc, "穿搭日记微信小程序", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "PRD v1.0", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "当前整理版本：基于已确认需求同步保存", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "代码目录：D:\\chuanda-project    产品文档目录：D:\\产品文档", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "1. 产品背景")
    add_paragraph(doc, "随着年轻女性对日常形象管理、穿搭表达和生活方式记录的关注度提升，“今天穿什么”已经成为高频、持续、带有情绪影响的日常决策问题。")
    add_paragraph(doc, "很多女生虽然拥有不少衣物，但仍然会在出门前遇到不知道怎么搭、衣橱利用率低、换季时难以判断穿着是否合适、缺少长期记录工具等问题。")
    add_paragraph(doc, "本产品希望以微信小程序作为主要载体，打造一款围绕“数字衣橱、今日穿搭决策、每日穿搭记录”展开的轻量工具型产品。")

    add_heading(doc, "2. 产品目标")
    add_bullets(doc, [
        "帮助用户更轻松地完成每日穿搭决策。",
        "提升衣橱使用效率，让用户更清楚自己有什么、哪些适合当前天气和场景。",
        "沉淀个人穿搭资产，通过长期记录形成可回看、可复用、可优化的数字衣橱。",
    ])
    add_paragraph(doc, "第一阶段成功标志：用户能够录入基础衣橱、获取天气与推荐建议、看到人物搭配预览、记录并回看穿搭。")

    add_heading(doc, "3. 用户画像")
    add_paragraph(doc, "核心用户：18-30 岁女性用户，关注日常穿搭体验，愿意为“更省心、更好看、更适合自己”的穿搭决策投入少量时间，但不希望操作过重。")
    add_bullets(doc, [
        "学生用户：更关注风格感、记录感和拍照表达。",
        "初入职场用户：更关注通勤效率、天气适配和不出错搭配。",
        "有记录习惯的穿搭爱好者：希望把自己的穿搭长期沉淀和复用。",
    ])

    add_heading(doc, "4. 用户痛点")
    add_bullets(doc, [
        "每天都可能陷入“今天穿什么”的高频决策焦虑。",
        "衣服很多，但并没有形成清晰的数字化整理，衣橱利用率不高。",
        "天气变化对穿搭影响很大，但很难快速判断怎么穿得舒服又好看。",
        "现有记录方式零散，不利于整理、搜索、复用和复盘。",
        "很多内容平台偏“看别人怎么穿”，缺少基于自己衣橱的实用建议。",
    ])

    add_heading(doc, "5. 核心使用场景")
    add_paragraph(doc, "5.1 出门前快速决策场景", bold=True)
    add_bullets(doc, [
        "用户希望在 1-3 分钟内得到今日穿搭建议。",
        "支持根据四季推荐穿衣风格。",
        "支持人物搭配模块，基于帽子、上衣、下装、鞋子组合生成虚拟人物穿搭预览。",
    ])
    add_paragraph(doc, "人物搭配模块策略：第一阶段采用“固定人物底图 + 单品叠加展示”，后续升级为“2D 虚拟模特换装系统”。")
    add_paragraph(doc, "5.2 衣橱整理场景", bold=True)
    add_bullets(doc, [
        "用户在空闲时间整理衣物，形成数字衣橱。",
        "支持录入扩展信息：购入渠道、购入时间、购入价格、品牌名称、尺码、喜好程度、已穿次数。",
    ])
    add_paragraph(doc, "5.3 穿搭记录与回看场景", bold=True)
    add_bullets(doc, [
        "支持保存每日穿搭，并能回看历史记录。",
        "长期形成可复用的穿搭资产库。",
    ])
    add_paragraph(doc, "5.4 天气辅助决策场景", bold=True)
    add_bullets(doc, [
        "支持自动获取天气、手动选择城市、手动调整温度。",
        "根据城市、天气和温度给出穿衣建议。",
    ])

    add_heading(doc, "6. 产品定位与价值主张")
    add_paragraph(doc, "产品定位：一款面向年轻女性用户、围绕“数字衣橱、今日穿搭推荐、每日穿搭记录”三件事展开的微信小程序。")
    add_paragraph(doc, "核心价值主张：用你自己的衣橱，为你今天的穿搭做决定。")
    add_bullets(doc, [
        "省心：减少“今天穿什么”的决策负担。",
        "实用：结合天气、场景和已有衣物给出能真正落地的建议。",
        "沉淀：把穿搭积累成自己的数字衣橱和穿搭记录。",
    ])

    add_heading(doc, "7. MVP 范围与优先级")
    add_paragraph(doc, "第一阶段目标：先做成一个可在微信小程序中真实使用、围绕“今日穿搭决策”展开的工具型产品。")
    add_bullets(doc, [
        "优先保证主链路闭环：查看天气 -> 获取推荐 -> 查看人物搭配 -> 记录今日穿搭。",
        "优先做高频功能和对推荐质量有直接帮助的功能。",
        "高复杂度但非第一阶段必要功能延后到第二版。",
    ])

    add_heading(doc, "8. 第一阶段核心功能清单")
    add_bullets(doc, [
        "首页：展示天气、季节风格入口、今日推荐入口、快捷记录入口。",
        "数字衣橱管理：上衣 / 下装 / 鞋子 / 帽子分类，支持新增、编辑、删除、查看详情。",
        "衣橱数据看板：总件数、上衣、下装、鞋子、帽子数量及品类分布。",
        "今日推荐模块：根据天气和季节生成多套搭配建议。",
        "人物搭配模块：固定人物底图 + 单品叠加展示。",
        "每日穿搭记录：保存、回看、查看详情。",
        "天气辅助模块：支持 6 天展示、切换城市、手动温度调整。",
        "AI 标签分析：上传单件衣物图后自动生成结构化标签，并支持手动修改。",
        "同类衣物对比：衣橱内同类单品对比。",
        "主题模式切换：浅色 / 深色 / 跟随系统。",
    ])

    add_heading(doc, "9. 功能删改说明")
    add_bullets(doc, [
        "删除：AI 抠图、智能分区入柜，后续不再设计和实现。",
        "新增：上传图片后 AI 自动分析衣物标签。",
        "新增：同类衣物对比功能。",
    ])

    add_heading(doc, "10. AI 标签分析模块")
    add_paragraph(doc, "功能定位：基于用户上传的单件衣物图片，自动生成衣物标签，并辅助用户完成衣橱录入。")
    add_bullets(doc, [
        "第一版仅支持单件衣物图。",
        "第一版以标签自动生成和入库为主，不强制上线完整分析报告页。",
        "第二版再增加完整分析报告页、身型建议、风格雷达图。",
        "AI 分析结果必须支持手动修改。",
    ])
    add_paragraph(doc, "第一版建议识别字段：")
    add_bullets(doc, [
        "基础信息：大类、品类、适合性别。",
        "颜色信息：主色、次色、强调色、主色明度、主色色度。",
        "适用信息：季节适用性、适用场合、适合人群。",
        "设计信息：服装长度、合身度、材料、图案、细节特征。",
        "品类特征：如上衣领型、袖长；下装裤型、裤长；帽型；鞋型等。",
        "搭配辅助信息：风格标签、搭配关键词、建议搭配颜色、搭配难度、搭配禁忌或提示。",
    ])

    add_heading(doc, "11. 衣橱录入与管理")
    add_paragraph(doc, "第一版衣橱录入主流程：上传单件衣物图片 -> AI 自动分析标签 -> 用户确认 / 修改结果 -> 补充字段 -> 保存到衣橱。")
    add_paragraph(doc, "第一版支持分类：上衣、下装、鞋子、帽子。")
    add_paragraph(doc, "第二版预留扩展品类：外套、包包、配饰。")
    add_paragraph(doc, "字段结构：")
    add_bullets(doc, [
        "必填：分类、名称、图片。",
        "选填：备注、购入渠道、购入时间、购入价格、品牌名称、尺码、喜好程度、已穿次数。",
        "AI 自动字段：风格标签、季节适用性、适用场合、颜色、材料、设计细节等。",
    ])

    add_heading(doc, "12. 服装对比功能")
    add_paragraph(doc, "第一版优先支持：衣橱内已上传衣物之间的同类对比。")
    add_bullets(doc, [
        "仅支持同一类衣物对比：上衣 vs 上衣、下装 vs 下装、鞋子 vs 鞋子、帽子 vs 帽子。",
        "对比维度：基础信息、颜色、适用信息、风格标签等。",
        "长期目标支持三种场景：衣橱内对比、已上传 vs 新上传、两件新上传待分析衣物对比。",
        "第一版优先做衣橱内对比。",
    ])

    add_heading(doc, "13. 主题模式切换")
    add_bullets(doc, [
        "第一版支持浅色模式、深色模式、跟随系统。",
        "支持用户在设置页切换，并保存设置。",
        "主题切换覆盖首页、衣橱页、推荐页、详情页、对比页、设置页及核心组件。",
    ])

    add_heading(doc, "14. 天气模块与首页修订")
    add_paragraph(doc, "第一版首页天气区域采用固定 6 天展示方案，即“今天 + 后 5 天”。")
    add_bullets(doc, [
        "默认只展示当天的日期、天气、温度区间与详细穿衣建议。",
        "天气趋势条保留 6 天日期入口。",
        "点击某一天日期后，天气主卡整卡切换为该日对应信息。",
        "每一天显示：日期、天气、温度区间、一句简短建议。",
        "当天显示详细建议，未来 5 天显示简短建议。",
    ])
    add_paragraph(doc, "首页四季风格推荐入口仅展示当前季节标签，不展示四季全部标签。")

    add_heading(doc, "15. 第二版规划")
    add_bullets(doc, [
        "完整分析报告页。",
        "身型与身材建议。",
        "风格雷达图。",
        "穿着次数统计。",
        "最久没穿分析。",
        "按月风格趋势图。",
        "包包 / 配饰扩展。",
        "更复杂的对比能力。",
        "2D 虚拟模特换装系统。",
    ])

    add_heading(doc, "16. 页面结构")
    add_paragraph(doc, "第一版建议采用 4 个底部主导航页 + 1 个中间快捷新增入口：")
    add_bullets(doc, [
        "首页",
        "衣橱",
        "添加",
        "搭配",
        "我的",
    ])
    add_paragraph(doc, "页面清单：")
    add_bullets(doc, [
        "首页",
        "衣橱页",
        "衣物分类列表页",
        "新增衣物页",
        "衣物详情页",
        "衣物对比页",
        "搭配推荐页",
        "人物搭配预览页",
        "记录穿搭页",
        "穿搭历史页",
        "穿搭详情页",
        "我的页",
        "设置页",
    ])

    add_heading(doc, "17. 核心页面模块布局草案")
    add_paragraph(doc, "17.1 首页", bold=True)
    add_bullets(doc, [
        "顶部基础区：标题、日期、问候语。",
        "今日天气主卡：当前城市、切换城市、当天天气、温度区间、详细穿衣建议。",
        "六天天气趋势条：默认选中今天，点击日期后整卡切换对应日天气信息。",
        "当季风格标签：仅展示当前季节风格入口。",
        "今日推荐入口卡。",
        "快捷记录入口。",
    ])
    add_paragraph(doc, "17.2 衣橱页", bold=True)
    add_bullets(doc, [
        "顶部标题区：页面标题、新增衣物按钮。",
        "数据看板区：总件数、各品类数量、分布图。",
        "分类导航区：上衣、下装、鞋子、帽子。",
        "最近新增 / 最近编辑区。",
    ])
    add_paragraph(doc, "17.3 衣物分类列表页", bold=True)
    add_bullets(doc, [
        "顶部标题区：分类名称、数量、右上角“筛选”“对比”文字按钮。",
        "筛选面板由顶部按钮触发，不常驻页面。",
        "筛选内容：搜索、部位筛选、颜色筛选、风格筛选、季节筛选，支持关闭 / 重置 / 应用。",
        "删除排序功能。",
        "衣物卡片列表区：图片、名称、核心标签、详情入口、对比模式勾选状态。",
        "对比模式下最多只能选择两件，同样不再使用底部操作条。",
    ])
    add_paragraph(doc, "17.4 新增衣物页", bold=True)
    add_bullets(doc, [
        "上传图片区。",
        "AI 分析状态区。",
        "AI 标签结果区。",
        "用户补充字段区。",
        "保存操作区。",
    ])
    add_paragraph(doc, "17.5 搭配推荐页", bold=True)
    add_bullets(doc, [
        "顶部天气摘要区为精简版，用于说明推荐依据，不重复首页完整天气模块。",
        "筛选条件区。",
        "推荐结果区。",
        "换一组推荐按钮。",
    ])
    add_paragraph(doc, "17.6 人物搭配预览页", bold=True)
    add_bullets(doc, [
        "搭配标题区。",
        "人物展示主区域：固定人物底图 + 单品叠加。",
        "单品信息区。",
        "搭配说明区。",
        "操作按钮区：换一套、记录今日穿搭、返回推荐列表。",
    ])
    add_paragraph(doc, "17.7 记录穿搭页", bold=True)
    add_bullets(doc, [
        "日期区。",
        "搭配结果区。",
        "备注区。",
        "保存区。",
    ])
    add_paragraph(doc, "17.8 我的页与设置页", bold=True)
    add_bullets(doc, [
        "我的页承载穿搭历史、设置、主题模式入口。",
        "设置页承载浅色、深色、跟随系统等主题模式切换。",
    ])

    add_heading(doc, "18. 已确认交互规则")
    add_bullets(doc, [
        "首页天气主卡默认显示当天信息，点击未来日期后整卡切换。",
        "未来 5 天只显示简短建议，当天显示详细建议。",
        "分类列表页顶部右上角仅保留“筛选”和“对比”，第一版不做“编辑”。",
        "对比模式下最多只能选择两件同类衣物。",
        "搭配推荐页仅显示精简天气摘要，不重复首页完整天气模块。",
    ])

    add_heading(doc, "19. 当前版本结论")
    add_paragraph(doc, "截至目前，第一版微信小程序的产品边界已经清晰：以天气驱动的今日穿搭决策为主线，以数字衣橱和 AI 标签分析为基础，以人物搭配展示和每日穿搭记录作为核心体验闭环。")
    add_paragraph(doc, "备注：本文件为当前沟通确认版，后续 PRD 新增或调整后需继续同步更新到同一目录。")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(str(OUTPUT_FILE))


if __name__ == "__main__":
    build_document()
