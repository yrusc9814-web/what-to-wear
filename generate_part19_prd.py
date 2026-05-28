from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


OUT = Path(r"D:\产品文档\穿搭日记微信小程序_PRD_第19部分_开发实施计划_v1.4_审核稿.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        set_cell_shading(hdr[i], "EDEDED")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.7)
    r = p.add_run(text)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10.5)


def style_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)

    for name, size, color in [
        ("Heading 1", 16, "111111"),
        ("Heading 2", 13, "222222"),
        ("Heading 3", 11, "333333"),
    ]:
        st = styles[name]
        st.font.name = "微软雅黑"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)


def build_doc():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("穿搭日记微信小程序 PRD\n第19部分：开发实施计划与里程碑")
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("v1.4 审核稿｜仅供审核，暂不替换正式PRD")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_heading("19.1 本节定位", level=1)
    doc.add_paragraph(
        "第19部分用于把已确认的产品范围、核心流程、数据结构和技术架构，转换为第一版开发实施计划。"
        "本节不新增产品功能，重点明确开发顺序、优先级、外部依赖、验收标准和启动清单。"
    )

    doc.add_heading("19.2 开发目标", level=1)
    add_bullet(doc, "完成第一版主链路闭环：衣物入库 -> 天气获取 -> 搭配推荐 -> 人物搭配预览 -> 穿搭记录 -> 衣橱沉淀。")
    add_bullet(doc, "确保第一版优先完成高频刚需能力，避免在未跑通主链路前投入第二版统计、报告或复杂算法。")
    add_bullet(doc, "将天气API、通义千问API Key、微信云开发环境作为开发启动前置依赖管理。")
    add_bullet(doc, "形成可执行的阶段计划，便于后续按阶段开发、测试、验收和同步PRD。")

    doc.add_heading("19.3 开发阶段划分", level=1)
    add_table(
        doc,
        ["阶段", "主要任务", "完成标志"],
        [
            ("阶段0：准备与环境", "确认微信小程序AppID、云开发环境、天气API、通义千问API Key、基础仓库结构。", "可以在本地和微信开发者工具中启动项目。"),
            ("阶段1：项目基础框架", "搭建小程序页面结构、底部导航、全局主题配置、云函数目录、基础工具方法。", "首页、衣橱、添加、搭配、我的五个入口可正常跳转。"),
            ("阶段2：衣橱基础能力", "实现衣物数据模型、图片上传、衣物列表、衣物详情、编辑保存、删除。", "用户可完成无AI参与的基础衣物入库与管理。"),
            ("阶段3：天气与首页", "接入天气云函数，支持城市选择、6天天气数据、今日详细建议、未来日期整卡切换。", "首页可稳定展示天气与穿衣建议，天气失败时保持获取中并支持切换城市。"),
            ("阶段4：AI标签分析", "接入通义千问视觉模型，上传单件衣物图后生成标签，支持用户修改后保存。", "AI成功时自动填充标签；超时或失败时可重试或取消。"),
            ("阶段5：搭配推荐与人物预览", "基于本地衣橱数据和天气信息生成推荐，展示组合结果和人物搭配预览。", "用户可从推荐结果进入记录穿搭页，且推荐数据能被带入。"),
            ("阶段6：穿搭记录、对比与主题", "实现穿搭记录保存、同类衣物对比、衣橱数据看板、深色/浅色/跟随系统主题。", "第一版关键辅助能力完成，并可进入整体联调。"),
            ("阶段7：联调测试与验收", "覆盖主链路、异常流程、权限流程、接口失败、数据保存、真机体验测试。", "满足第一版交付标准，可进入小范围试用或提审准备。"),
        ],
    )

    doc.add_heading("19.4 功能优先级", level=1)
    add_table(
        doc,
        ["优先级", "范围", "说明"],
        [
            ("P0", "项目基础、衣物入库、AI标签分析、天气获取、搭配推荐、人物搭配预览、穿搭记录、主题模式。", "第一版主链路必需能力；缺少任一项都会影响产品闭环。"),
            ("P1", "衣橱数据看板、同类衣物对比、筛选、异常提示完善、基础真机适配。", "第一版上线前应尽量完成的增强能力；不应反向阻塞P0主链路。"),
            ("P2", "风格趋势、最久没穿搭配、身型建议、风格雷达图、外套/包包/配饰、复杂AI分析报告。", "第二版或后续版本能力，第一版仅保留数据和结构上的扩展空间。"),
        ],
    )

    doc.add_heading("19.5 外部服务接入计划", level=1)
    add_table(
        doc,
        ["服务", "用途", "当前状态", "处理原则"],
        [
            ("天气API", "根据城市获取当天及未来5天天气、温度区间，并生成穿衣建议。", "待申请。", "通过云函数调用，不在前端暴露Key；温度由城市天气数据返回，用户不可手动修改温度。"),
            ("通义千问视觉模型", "识别单件衣物图，生成衣物标签和基础字段。", "待申请API Key。", "通过云函数调用，不在前端暴露Key；第一版只分析单件衣物图。"),
            ("微信云开发", "云数据库、云存储、云函数、用户openid识别。", "待创建或确认环境。", "统一使用 openid 作为用户唯一标识；云函数从上下文获取 openid，不信任前端传入的 userId。"),
        ],
    )

    doc.add_heading("19.6 第一版开发依赖", level=1)
    add_bullet(doc, "天气服务商、天气Key、城市查询方式确认后，才能进入真实天气接口联调。")
    add_bullet(doc, "通义千问API Key确认后，才能进入AI识别云函数联调；在此之前可先使用Mock数据开发前端流程。")
    add_bullet(doc, "微信小程序AppID和云开发环境确认后，才能正式接入云数据库、云存储和云函数。")
    add_bullet(doc, "人物搭配预览第一版采用固定人物底图+单品图片叠层方案，不引入AI抠图。")
    add_bullet(doc, "推荐逻辑第一版放在前端本地服务中实现，后续复杂推荐再迁移到云函数。")

    doc.add_heading("19.7 测试与验收标准", level=1)
    add_table(
        doc,
        ["验收项", "通过标准"],
        [
            ("衣物录入", "用户可选择图片、触发AI分析、修改标签、补充选填字段并保存到衣橱。"),
            ("权限处理", "用户拒绝上传媒体权限后终止本次新增衣服流程，不进入无图片手动录入。"),
            ("AI异常", "AI超时或失败时，用户可选择重新分析或取消，不允许跳过AI后继续保存本次图片录入。"),
            ("天气展示", "首页默认展示当天完整天气卡；点击未来日期后整卡切换为对应日期信息。"),
            ("天气异常", "天气获取失败时持续保持获取中，支持用户手动切换城市，推荐入口在成功前禁用或弱化。"),
            ("搭配推荐", "推荐结果至少包含上衣、下装、鞋子、帽子中的可用组合，并能进入人物预览。"),
            ("人物预览", "用户能看到固定人物底图上的搭配叠层预览，且可进入记录穿搭页。"),
            ("穿搭记录", "同一日期允许新增多条穿搭记录，不覆盖已有记录。"),
            ("对比功能", "用户可在同类衣物中最多选择两件进行对比。"),
            ("主题模式", "支持深色模式、浅色模式、跟随系统设备三种设置。"),
        ],
    )

    doc.add_heading("19.8 风险与处理策略", level=1)
    add_table(
        doc,
        ["风险", "影响", "处理策略"],
        [
            ("天气API申请或联调延迟", "首页天气和推荐入口无法进入真实数据联调。", "先使用Mock天气数据开发页面和推荐流程，Key到位后替换为真实接口。"),
            ("通义千问识别结果不稳定", "衣物标签质量影响推荐准确度。", "第一版允许用户手动修改AI结果，并限制AI输出为结构化字段。"),
            ("图片格式差异", "微信上传图片可能包含jpg、png、heic等格式。", "存储路径文件后缀按实际格式动态获取，不硬编码jpg。"),
            ("人物叠层效果不理想", "搭配预览观感影响核心体验。", "第一版先做规则化层级和位置，后续根据真实素材优化叠层模板。"),
            ("推荐结果不足", "衣橱单品过少时无法形成完整搭配。", "给出缺少品类提示，并允许用户继续添加对应品类衣物。"),
        ],
    )

    doc.add_heading("19.9 开发启动清单", level=1)
    add_bullet(doc, "确认微信小程序AppID。")
    add_bullet(doc, "创建或确认微信云开发环境。")
    add_bullet(doc, "申请天气API并记录服务商、Key、城市查询接口。")
    add_bullet(doc, "申请通义千问API Key，并确认使用DashScope/百炼控制台入口。")
    add_bullet(doc, "确认第一版页面清单、底部导航、添加快捷菜单和主流程入口。")
    add_bullet(doc, "确认云数据库集合：users、clothing_items、outfit_records。")
    add_bullet(doc, "确认图片云存储目录和动态文件后缀策略。")
    add_bullet(doc, "准备Mock数据：天气、AI识别结果、衣物样例、推荐结果。")

    doc.add_heading("19.10 与后续章节的关系", level=1)
    doc.add_paragraph(
        "第19部分确认后，后续可进入第20部分接口与字段说明。第20部分建议聚焦云函数入参/出参、天气数据归一化结构、"
        "AI分析返回JSON结构、前端服务方法和错误码。第20部分完成后，即可开始正式搭建微信小程序项目。"
    )

    doc.save(OUT)


if __name__ == "__main__":
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build_doc()
    print(OUT)
