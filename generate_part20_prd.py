from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pathlib import Path


OUT = Path(r"D:\产品文档\穿搭日记微信小程序_PRD_第20部分_接口与字段说明_v1.4_审核稿.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "EDEDED")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.2)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F6F6F6")
    p._p.get_or_add_pPr().append(shading)


def style_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        st = doc.styles[name]
        st.font.name = "微软雅黑"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(30, 30, 30)


def build():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("穿搭日记微信小程序 PRD\n第20部分：接口与字段说明")
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("v1.4 审核稿｜仅供审核，暂不替换正式PRD")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_heading("20.1 本节定位", level=1)
    doc.add_paragraph(
        "第20部分用于定义第一版开发所需的云函数、前端服务、核心入参/出参、统一错误码和Mock数据要求。"
        "本节不新增产品功能，只把第17部分数据结构、第18部分技术架构和第19部分实施计划落到可联调的接口层。"
    )

    doc.add_heading("20.2 接口设计原则", level=1)
    add_bullet(doc, "用户身份统一使用 openid，云函数必须从微信云开发上下文获取 openid，不信任前端传入的 userId/openid。")
    add_bullet(doc, "天气API和通义千问API Key只允许配置在云函数环境变量或云端安全配置中，不允许出现在小程序前端代码。")
    add_bullet(doc, "云函数返回结构保持统一：success、code、message、data。前端只依赖统一结构，不直接依赖第三方原始返回。")
    add_bullet(doc, "第一版推荐逻辑放在前端本地服务，不作为云函数；后续复杂推荐再迁移到云函数。")
    add_bullet(doc, "第一版对比功能为前端临时结构，不持久化数据库。")

    doc.add_heading("20.3 云函数与前端服务清单", level=1)
    doc.add_paragraph(
        "说明：第18部分已定义第一版关键云函数。本节为了支撑完整开发联调，对衣物CRUD和记录查询云函数进行补充。"
        "正式合并时需同步修订第18部分云函数清单，避免章节之间出现实现范围不一致。"
    )
    add_table(
        doc,
        ["名称", "类型", "用途", "第一版是否必须"],
        [
            ("getWeather", "云函数", "根据城市获取当天及未来5天天气，并返回归一化天气结构。", "是"),
            ("analyzeClothing", "云函数", "调用通义千问视觉模型，分析单件衣物图片并返回结构化标签。", "是"),
            ("saveClothing", "云函数", "保存衣物基础信息、AI标签、选填字段和图片fileID。", "是"),
            ("getWardrobe", "云函数", "查询当前用户衣橱列表，支持分类和筛选条件。", "是"),
            ("updateClothing", "云函数", "更新衣物信息和用户手动修正后的标签。", "是"),
            ("deleteClothing", "云函数", "删除或软删除衣物记录。", "是"),
            ("saveOutfitRecord", "云函数", "保存每日穿搭记录，允许同一日期多条记录。", "是"),
            ("getOutfitRecords", "云函数", "查询穿搭记录列表，为后续统计预留。", "是"),
            ("recommendationService", "前端本地服务", "基于衣橱、天气、场景和季节生成搭配推荐。", "是"),
            ("compareClothing", "前端临时结构", "同类衣物最多两件对比，不持久化。", "是"),
        ],
    )

    doc.add_heading("20.4 统一返回结构", level=1)
    add_code_block(
        doc,
        "{\n"
        '  "success": true,\n'
        '  "code": "OK",\n'
        '  "message": "success",\n'
        '  "data": {}\n'
        "}",
    )
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ("success", "boolean", "本次请求是否成功。"),
            ("code", "string", "统一状态码。成功为OK，失败使用第20.13节错误码。"),
            ("message", "string", "给前端或调试使用的简短说明。"),
            ("data", "object/null", "业务数据。失败时可为空。"),
        ],
    )

    doc.add_heading("20.5 getWeather 天气接口", level=1)
    doc.add_paragraph("用途：根据用户选择的城市获取当天及未来5天的天气、温度区间和穿衣建议。温度不允许用户手动修改。")
    add_table(
        doc,
        ["入参", "类型", "必填", "说明"],
        [
            ("cityName", "string", "是", "用户当前选择城市，如天津市。"),
            ("cityCode", "string", "否", "天气服务商返回的城市ID。首次可为空，后续建议缓存。"),
            ("days", "number", "否", "默认6，表示当天+未来5天。"),
        ],
    )
    add_table(
        doc,
        ["返回字段", "类型", "说明"],
        [
            ("cityName", "string", "展示城市名称。"),
            ("updatedAt", "string", "天气数据更新时间。"),
            ("daily", "array", "长度为6的天气数组。"),
            ("daily[].date", "string", "日期，格式YYYY-MM-DD。"),
            ("daily[].weatherText", "string", "天气描述，如晴、多云、小雨。"),
            ("daily[].tempMin", "number", "最低温。"),
            ("daily[].tempMax", "number", "最高温。"),
            ("daily[].suggestionShort", "string", "每日一句简短穿衣建议。"),
            ("daily[].suggestionDetail", "string", "当天详细建议；未来日期可为空或短文案。"),
        ],
    )
    doc.add_paragraph("异常规则：天气获取失败时前端保持获取中状态，支持用户切换城市；推荐入口在成功前禁用或弱化。云函数可按固定间隔重试，但不得高频连续请求第三方接口。")

    doc.add_heading("20.6 analyzeClothing AI衣物分析接口", level=1)
    doc.add_paragraph("用途：对单件衣物图片进行AI分析，返回可编辑的衣物标签。第一版不支持多图分析，不支持AI抠图。")
    add_table(
        doc,
        ["入参", "类型", "必填", "说明"],
        [
            ("fileID", "string", "是", "微信云存储图片fileID。"),
            ("categoryHint", "string", "否", "用户预选大类，如top/bottom/shoes/hat。"),
            ("imageFormat", "string", "否", "图片格式，如jpg/png/heic；后缀按实际格式动态获取。"),
        ],
    )
    add_table(
        doc,
        ["返回字段", "类型", "说明"],
        [
            ("aiAnalysisStatus", "string", "pending/success/failed。"),
            ("category", "string", "衣物大类：top/bottom/shoes/hat。"),
            ("subCategory", "string", "品类，如针织衫、阔腿裤、运动鞋。"),
            ("name", "string", "AI建议图片名称，用户可修改。"),
            ("colors", "array", "主色、次色、强调色等。"),
            ("styleTags", "array", "风格标签，如复古、学院、街头。"),
            ("seasonTags", "array", "适用季节：spring/summer/autumn/winter。"),
            ("occasionTags", "array", "适用场景，如日常、约会、通勤。"),
            ("material", "string", "材质识别结果，可为空。"),
            ("pattern", "string", "图案识别结果，可为空。"),
            ("confidence", "number", "AI识别置信度，0到1。"),
        ],
    )
    doc.add_paragraph("异常规则：请求超过10秒未返回视为超时。超时或失败时，用户可重新分析或取消；不允许跳过AI后继续保存本次图片录入。")

    doc.add_heading("20.7 saveClothing 衣物保存接口", level=1)
    add_table(
        doc,
        ["入参", "类型", "必填", "说明"],
        [
            ("fileID", "string", "是", "衣物图片云存储ID。"),
            ("category", "string", "是", "top/bottom/shoes/hat。"),
            ("subCategory", "string", "否", "具体品类，如针织衫、阔腿裤；AI未识别时可为空，用户可后续补充。"),
            ("name", "string", "是", "衣物名称。"),
            ("colors", "array", "否", "颜色标签。"),
            ("styleTags", "array", "否", "风格标签。"),
            ("seasonTags", "array", "否", "季节标签。"),
            ("occasionTags", "array", "否", "场景标签。"),
            ("optionalFields", "object", "否", "购入渠道、时间、价格、品牌、尺码、喜好程度、已穿次数等选填字段。"),
        ],
    )
    doc.add_paragraph("保存规则：openid由云函数上下文获取；createdAt、updatedAt由服务端生成；用户可手动修改AI分析结果后保存。")

    doc.add_heading("20.8 updateClothing 衣物更新接口", level=1)
    add_table(
        doc,
        ["入参", "类型", "必填", "说明"],
        [
            ("clothingId", "string", "是", "需要更新的衣物ID。"),
            ("patch", "object", "是", "需要更新的字段集合。"),
            ("patch.name", "string", "否", "衣物名称。"),
            ("patch.category", "string", "否", "衣物大类：top/bottom/shoes/hat。"),
            ("patch.subCategory", "string", "否", "具体品类。"),
            ("patch.colors", "array", "否", "颜色标签。"),
            ("patch.styleTags", "array", "否", "风格标签。"),
            ("patch.seasonTags", "array", "否", "季节标签。"),
            ("patch.occasionTags", "array", "否", "场景标签。"),
            ("patch.optionalFields", "object", "否", "购入渠道、时间、价格、品牌、尺码、喜好程度、已穿次数等选填字段。"),
        ],
    )
    doc.add_paragraph(
        "校验规则：只能更新当前 openid 名下的衣物；clothingId 不存在或不属于当前用户时返回权限或不存在错误。"
        "updatedAt 由服务端更新。"
    )

    doc.add_heading("20.9 deleteClothing 衣物删除接口", level=1)
    add_table(
        doc,
        ["入参", "类型", "必填", "说明"],
        [
            ("clothingId", "string", "是", "需要删除的衣物ID。"),
        ],
    )
    doc.add_paragraph(
        "删除规则：第一版建议采用软删除，设置 isDeleted=true 和 deletedAt，不直接物理删除图片和记录。"
        "若衣物已被历史穿搭记录引用，前端展示历史记录时应检查各单品 isDeleted 状态："
        "isDeleted=true 的单品在该位置显示灰色占位块 + “该单品已删除”文字提示，不隐藏该位置，"
        "确保历史穿搭记录完整性，并为第二版穿着次数统计保留完整数据基础。"
    )
    doc.add_paragraph(
        "交互提示：用户删除衣物前需进行二次确认，建议文案为“删除后，该衣物将从衣橱中隐藏，"
        "但历史穿搭记录中会显示为‘该单品已删除’。”"
    )

    doc.add_heading("20.10 getWardrobe 衣橱查询接口", level=1)
    add_table(
        doc,
        ["入参", "类型", "必填", "说明"],
        [
            ("category", "string", "否", "按品类筛选。"),
            ("colorTags", "array", "否", "按颜色筛选。"),
            ("styleTags", "array", "否", "按风格筛选。"),
            ("seasonTags", "array", "否", "按季节筛选。"),
            ("keyword", "string", "否", "按名称或品牌搜索。"),
            ("pageSize", "number", "否", "默认50，最大100。"),
            ("pageToken", "string", "否", "下一页标识；第一版可为空。"),
        ],
    )
    add_table(
        doc,
        ["返回字段", "类型", "说明"],
        [
            ("items", "array", "衣物列表。"),
            ("nextPageToken", "string/null", "下一页标识，无更多数据时为空。"),
            ("dashboard.totalCount", "number", "衣橱总件数。"),
            ("dashboard.categoryCounts", "object", "上衣、下装、鞋子、帽子数量分布。"),
        ],
    )

    doc.add_paragraph("分页规则：第一版可先按默认pageSize返回，保留pageToken字段用于后续衣橱数据量变大后的平滑扩展。")

    doc.add_heading("20.11 recommendationService 本地推荐服务", level=1)
    doc.add_paragraph("第一版推荐逻辑在前端本地实现，不调用云函数。输入为衣橱列表、天气数据、日期、场景和季节标签，输出为临时RecommendationResult。")
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ("weatherDaily", "object", "当前选中日期的天气对象。"),
            ("wardrobeItems", "array", "当前用户衣橱列表。"),
            ("selectedDate", "string", "推荐日期。"),
            ("occasion", "string", "用户选择场景，可为空。"),
            ("result.topId", "string", "推荐上衣ID，可为空。"),
            ("result.bottomId", "string", "推荐下装ID，可为空。"),
            ("result.shoesId", "string", "推荐鞋子ID，可为空。"),
            ("result.hatId", "string", "推荐帽子ID，可为空。"),
            ("result.reason", "string", "推荐理由。"),
        ],
    )
    doc.add_paragraph("规则：如果衣橱单品不足以生成完整搭配，应返回缺少品类提示，并引导用户添加对应品类。")

    doc.add_heading("20.12 saveOutfitRecord 穿搭记录接口", level=1)
    add_table(
        doc,
        ["入参", "类型", "必填", "说明"],
        [
            ("date", "string", "是", "穿搭日期，格式YYYY-MM-DD。"),
            ("topId", "string", "否", "上衣ID。"),
            ("bottomId", "string", "否", "下装ID。"),
            ("shoesId", "string", "否", "鞋子ID。"),
            ("hatId", "string", "否", "帽子ID。"),
            ("weatherSnapshot", "object", "否", "记录当日天气快照。"),
            ("source", "string", "是", "记录来源：manual/recommendation。"),
            ("note", "string", "否", "用户备注。"),
        ],
    )
    doc.add_paragraph("校验规则：保存穿搭记录时，topId/bottomId/shoesId/hatId至少有一个字段有值。同一日期允许保存多条记录，不覆盖已有记录。")

    doc.add_heading("20.13 getOutfitRecords 穿搭记录查询接口", level=1)
    add_table(
        doc,
        ["入参", "类型", "必填", "说明"],
        [
            ("date", "string", "否", "按单日查询，格式YYYY-MM-DD。"),
            ("startDate", "string", "否", "按日期范围查询的开始日期。"),
            ("endDate", "string", "否", "按日期范围查询的结束日期。"),
            ("pageSize", "number", "否", "默认50，最大100。"),
            ("pageToken", "string", "否", "下一页标识；第一版可为空。"),
        ],
    )
    add_table(
        doc,
        ["返回字段", "类型", "说明"],
        [
            ("records", "array", "穿搭记录列表。"),
            ("nextPageToken", "string/null", "下一页标识，无更多数据时为空。"),
        ],
    )
    doc.add_paragraph("查询规则：只返回当前 openid 名下记录；同一日期存在多条记录时按创建时间倒序返回。")

    doc.add_heading("20.14 compareClothing 前端对比结构", level=1)
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ("selectedIds", "array", "最多两个衣物ID。"),
            ("category", "string", "必须为同一衣物大类。"),
            ("compareFields", "array", "名称、品类、颜色、风格、季节、适用场景、品牌、尺码等。"),
        ],
    )
    doc.add_paragraph("规则：第一版仅支持已上传衣物之间的同类对比；最多选择两件；对比结果不持久化。")

    doc.add_heading("20.15 图片上传与存储字段", level=1)
    add_table(
        doc,
        ["字段", "说明"],
        [
            ("fileID", "微信云存储返回的文件ID，数据库保存该字段。"),
            ("cloudPath", "云存储路径，建议包含openid、时间戳和动态后缀。"),
            ("ext", "根据实际图片格式动态获取，不硬编码jpg。"),
            ("size", "图片大小，后续可用于压缩策略。"),
        ],
    )
    doc.add_paragraph("权限规则：用户拒绝上传媒体权限后，终止本次新增衣物流程，不进入无图片手动录入。")

    doc.add_heading("20.16 统一错误码", level=1)
    add_table(
        doc,
        ["错误码", "含义", "前端处理"],
        [
            ("OK", "请求成功。", "正常展示结果。"),
            ("AUTH_OPENID_MISSING", "无法获取openid。", "提示用户重新进入小程序或重试。"),
            ("MEDIA_PERMISSION_DENIED", "用户拒绝上传媒体权限。", "终止新增衣服流程。"),
            ("VALIDATION_ERROR", "入参不合法。", "提示用户补全或修正信息。"),
            ("RESOURCE_NOT_FOUND", "数据不存在。", "提示记录不存在或已删除。"),
            ("RESOURCE_FORBIDDEN", "数据不属于当前用户。", "阻止访问或修改。"),
            ("AI_TIMEOUT", "AI分析超过10秒。", "展示重新分析/取消。"),
            ("AI_FAILED", "AI分析失败。", "展示重新分析/取消。"),
            ("WEATHER_FETCHING", "天气仍在获取中。", "保持获取中，推荐入口禁用或弱化。"),
            ("WEATHER_FAILED", "天气接口请求失败。", "继续重试，允许切换城市。"),
            ("WARDROBE_EMPTY", "衣橱暂无可用衣物。", "引导用户新增衣物。"),
            ("RECOMMENDATION_NOT_ENOUGH_ITEMS", "缺少生成搭配所需品类。", "提示缺少的品类。"),
            ("COMPARE_CATEGORY_MISMATCH", "对比衣物不是同一类别。", "提示只能选择同类衣物。"),
            ("COMPARE_LIMIT_EXCEEDED", "对比选择超过两件。", "阻止继续选择。"),
        ],
    )

    doc.add_heading("20.17 Mock数据要求", level=1)
    add_bullet(doc, "天气Mock：至少包含6天数据，每天包含日期、天气、最低温、最高温、简短建议；当天额外包含详细建议。")
    add_bullet(doc, "AI分析Mock：至少覆盖上衣、下装、鞋子、帽子四类，每类提供颜色、风格、季节、场景等标签。")
    add_bullet(doc, "衣橱Mock：至少准备每类3件衣物，用于测试推荐、筛选、对比和人物预览。")
    add_bullet(doc, "异常Mock：覆盖AI超时、AI失败、天气获取中、天气失败、权限拒绝、衣橱为空、推荐单品不足。")

    doc.add_heading("20.18 开发联调顺序", level=1)
    add_table(
        doc,
        ["顺序", "内容", "说明"],
        [
            ("1", "先用Mock数据跑通页面和主流程。", "避免天气API和AI Key未到位阻塞前端开发。"),
            ("2", "接入微信云开发和openid。", "先解决用户身份、数据库和云存储。"),
            ("3", "接入天气getWeather。", "完成首页天气和推荐入口状态联调。"),
            ("4", "接入analyzeClothing。", "完成图片上传、AI分析、用户修改、保存。"),
            ("5", "联调推荐和人物预览。", "验证推荐结果能传递到记录穿搭页。"),
            ("6", "补齐异常、对比、主题和真机测试。", "进入第一版整体验收。"),
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build()
    print(OUT)
