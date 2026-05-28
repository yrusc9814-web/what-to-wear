import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOC_DIR = Path(r"D:\chuanda-project\产品文档")
SRC = DOC_DIR / "穿搭日记微信小程序_PRD_v1.0.docx"
OUT = DOC_DIR / "穿搭日记微信小程序_PRD_v1.5_审核稿.docx"


def set_font(run, size=10.5, bold=False):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_font(r, 9, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell, fill="EDEDED"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table_after(anchor_p, headers, rows):
    doc = anchor_p._parent
    table = doc.add_table(rows=1, cols=len(headers), width=Cm(16))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        shade(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    anchor_p._p.addnext(table._tbl)
    return table


def remove_table(table):
    el = table._tbl
    el.getparent().remove(el)


def insert_paragraph_before(doc, index, text, style=None):
    p = doc.add_paragraph()
    if style:
        prefix = {"Heading 1": "1.", "Heading 2": "1.1", "Heading 3": "1.1.1"}.get(style)
        matched = None
        if prefix:
            for existing in doc.paragraphs:
                if existing.text.strip().startswith(prefix):
                    matched = existing.style
                    break
        if matched is not None:
            p.style = matched
    r = p.add_run(text)
    set_font(r, 10.5, style and style.startswith("Heading"))
    body = doc.element.body
    body.insert(index, p._p)
    return p


def add_styled_paragraph(doc, text, style=None):
    p = doc.add_paragraph()
    if style:
        prefix = {"Heading 1": "1.", "Heading 2": "1.1", "Heading 3": "1.1.1"}.get(style)
        matched = None
        if prefix:
            for existing in doc.paragraphs:
                if existing.text.strip().startswith(prefix):
                    matched = existing.style
                    break
        if matched is not None:
            p.style = matched
    r = p.add_run(text)
    set_font(r, 10.5, style and style.startswith("Heading"))
    return p


def body_index_of_paragraph(p):
    return list(p._parent._element).index(p._p)


def replace_in_runs(paragraph, replacements):
    text = paragraph.text
    new = text
    for old, rep in replacements.items():
        new = new.replace(old, rep)
    if new != text:
        for r in paragraph.runs:
            r.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new
        else:
            run = paragraph.add_run(new)
            set_font(run)


def replace_everywhere(doc, replacements):
    for p in doc.paragraphs:
        replace_in_runs(p, replacements)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_runs(p, replacements)


def find_heading(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def add_version_row(doc):
    table = doc.tables[0]
    exists = any(row.cells[0].text.strip() == "v1.5" for row in table.rows)
    if exists:
        return
    row = table.add_row().cells
    values = [
        "v1.5",
        "审核稿",
        "修正章节编号断层，补充隐私合规与冷启动流程；第一版搭配预览由人物叠图降级为单品平铺卡片；明确未来日期仅展示天气建议、当季风格标签、颜色协调策略与图片必填说明。",
    ]
    for i, v in enumerate(values):
        set_cell_text(row[i], v)


def insert_section_15_16(doc):
    anchor = find_heading(doc, "17.1")
    if not anchor:
        raise RuntimeError("Cannot find 17.1 anchor")
    idx = body_index_of_paragraph(anchor)
    items = [
        ("15. 隐私合规与个人信息保护", "Heading 1"),
        ("15.1 信息收集清单", "Heading 2"),
    ]
    # Rebuild ordered content by inserting before 17.1 one by one.
    insert_paragraph_before(doc, idx, "15. 隐私合规与个人信息保护", "Heading 1")
    idx += 1
    idx = body_index_of_paragraph(anchor)
    for text, style in [
        ("15.1 信息收集清单", "Heading 2"),
    ]:
        insert_paragraph_before(doc, idx, text, style)
        idx += 1
    tbl_anchor = doc.paragraphs[[p._p for p in doc.paragraphs].index(anchor._p) - 1]
    add_table_after(
        tbl_anchor,
        ["信息类型", "使用目的", "最小必要性说明", "是否可拒绝"],
        [
            ("微信 openid", "用户唯一身份标识", "必须，无法替代", "否"),
            ("位置信息", "自动获取当地天气", "可选，拒绝后手动选城市降级", "是"),
            ("相册/拍照", "上传衣物图片", "核心功能必须，拒绝后终止录入", "是"),
            ("衣物图片", "AI 分析标签、云端存储展示", "用户主动上传，第一版无图片不可保存", "是"),
            ("穿搭记录", "历史回看、第二版统计", "核心业务数据", "否"),
        ],
    )
    idx = body_index_of_paragraph(anchor)
    for text, style in [
        ("15.2 必须实现的合规动作", "Heading 2"),
        ("上线前必须提供隐私政策页面，说明信息收集范围、使用目的、保存期限、第三方服务（天气 API、通义千问 VL）、用户查阅/更正/删除权利和联系方式。", None),
        ("上线前必须接入微信官方隐私保护规范：在 app.json 中声明 __usePrivacyCheck__: true；用户触发需授权功能前，先调用 wx.getPrivacySetting 检查同意状态，未同意则拉起官方隐私授权页（wx.openPrivacyContract），用户拒绝后中止对应功能，不可绕过。自定义弹窗不可替代官方组件。", None),
        ("用户同意通过官方隐私授权页完成；同意时间通过 wx.onPrivacyAuthorizationSucceed 回调写入 privacyConsentedAt 字段，不再额外弹出自定义同意弹窗。", None),
        ("app.json 中还需声明位置权限用途：\"permission\": {\"scope.userLocation\": {\"desc\": \"用于获取当前城市天气，为你生成穿搭推荐\"}}。", None),
        ("位置权限仅在用户主动触发自动天气时申请，不在进入首页时预申请。", None),
        ("新增衣物页首次触发 AI 分析前，提示“图片将发送至 AI 服务进行分析”，用户确认后继续。", None),
        ("设置页需提供“注销账号/清除所有数据”入口，触发后软删除用户相关记录，硬删除可由后台定期处理。", None),
        ("15.3 User 对象新增合规字段", "Heading 2"),
    ]:
        insert_paragraph_before(doc, idx, text, style)
        idx += 1
    tbl_anchor = doc.paragraphs[[p._p for p in doc.paragraphs].index(anchor._p) - 1]
    add_table_after(
        tbl_anchor,
        ["字段", "含义", "说明"],
        [
            ("privacyConsentedAt", "用户同意隐私政策时间", "首次同意时写入，不可为空。"),
            ("privacyPolicyVersion", "同意时的隐私政策版本号", "隐私政策更新后可重新触发确认。"),
            ("dataDeleteRequestedAt", "用户申请删除数据时间", "用户触发注销或清除数据时写入。"),
        ],
    )
    idx = body_index_of_paragraph(anchor)
    for text, style in [
        ("16. 用户核心流程与新用户冷启动", "Heading 1"),
        ("16.1 新用户状态定义", "Heading 2"),
        ("新用户状态定义为：当前 openid 名下 clothing_items 有效记录数为 0。有效记录指 isDeleted=false 的衣物。", None),
        ("16.2 冷启动流程", "Heading 2"),
        ("首次进入且衣橱为空时展示引导页，文案为“用你自己的衣橱，决定今天穿什么”，主按钮“开始整理我的衣橱”，点击进入新增衣物页。引导页提供“跳过”次要按钮（右上角文字按钮），点击后进入首页空状态，不跳转衣橱页。", None),
        ("用户点击“跳过”后，记录 hasSkippedOnboarding=true；后续若衣橱仍为空，不再反复展示整页引导，只在首页与推荐页展示空状态引导。一旦用户成功保存第一件衣物，引导页永久不再展示。", None),
        ("首页空状态下，天气模块正常展示；今日推荐入口替换为引导卡“先添加几件衣物，推荐就能为你量身定制”，点击进入新增衣物页。", None),
        ("推荐页空状态下，不执行推荐逻辑，展示“衣橱还是空的，先去添加几件吧”与新增衣物按钮。", None),
        ("16.3 推荐可用性分层", "Heading 2"),
    ]:
        insert_paragraph_before(doc, idx, text, style)
        idx += 1
    tbl_anchor = doc.paragraphs[[p._p for p in doc.paragraphs].index(anchor._p) - 1]
    add_table_after(
        tbl_anchor,
        ["衣橱有效单品数", "推荐行为", "界面提示"],
        [
            ("0 件", "不生成推荐，显示空状态引导", "先添加衣物才能获取推荐"),
            ("1-3 件或缺某品类", "可生成部分推荐，缺失品类在 2×2 搭配预览卡片中显示灰色占位块 + “暂无该品类”文字，不隐藏格子，不改变卡片网格结构", "你的衣橱还缺 [品类]，添加后推荐更完整"),
            ("各核心品类均有", "正常推荐流程", "无提示"),
        ],
    )
    idx = body_index_of_paragraph(anchor)
    for text, style in [
        ("16.4 穿搭记录引导", "Heading 2"),
        ("保存穿搭记录时仍允许 topId/bottomId/shoesId/hatId 至少一个字段有值。若用户只选择 1 件单品，页面展示轻提示“当前只选了 1 件，搭配记录会更完整哦”，但不强制拦截。", None),
    ]:
        insert_paragraph_before(doc, idx, text, style)
        idx += 1


def insert_18_4(doc):
    anchor = find_heading(doc, "18.5")
    if not anchor:
        raise RuntimeError("Cannot find 18.5 anchor")
    idx = body_index_of_paragraph(anchor)
    for text, style in [
        ("18.4 前端技术选型说明", "Heading 1"),
        ("第一版采用微信小程序原生框架 + 微信云开发。页面样式使用全局 CSS 变量承载浅色、深色、跟随系统三种主题，避免后期主题返工。", None),
        ("推荐逻辑第一版在前端本地 recommendation-service.js 中实现；云函数只承担天气、AI、衣物 CRUD、穿搭记录等服务中转。", None),
        ("第一版搭配预览采用 2×2 单品平铺卡片，不做人物叠图，不依赖 AI 抠图。第二版再评估抠图/虚拟模特换装能力。", None),
    ]:
        insert_paragraph_before(doc, idx, text, style)
        idx += 1


def insert_12_9_12_10(doc):
    if any(p.text.strip().startswith("12.9 穿搭历史页") for p in doc.paragraphs):
        return
    anchor = find_heading(doc, "13.")
    if not anchor:
        return
    idx = body_index_of_paragraph(anchor)
    for text, style in [
        ("12.9 穿搭历史页", "Heading 2"),
        ("顶部标题区：页面标题“穿搭历史”。", None),
        ("列表区：按日期倒序排列，每条记录展示日期、单品缩略图组合（最多展示4个小图）、备注摘要（若有）；同一天有多条记录时逐条展示，不合并。", None),
        ("空状态：无任何记录时展示“还没有穿搭记录，去记录今天的穿搭吧”与快捷记录入口。", None),
        ("点击单条记录进入穿搭详情页。", None),
        ("12.10 穿搭详情页", "Heading 2"),
        ("顶部标题区：展示穿搭日期。", None),
        ("搭配单品区：展示该条记录的上衣、下装、鞋子、帽子图片与名称；若单品已被删除，对应位置显示灰色占位块 + “该单品已删除”。", None),
        ("天气快照区（若有）：展示记录时的城市、天气状态、温度区间。", None),
        ("备注区：展示用户备注，若无则不展示。", None),
        ("操作区：提供“删除记录”入口，触发二次确认后软删除；第一版不支持编辑已有穿搭记录。", None),
    ]:
        insert_paragraph_before(doc, idx, text, style)
        idx += 1


def insert_top_heading_if_missing(doc, before_prefix, heading_text):
    if any(p.text.strip() == heading_text for p in doc.paragraphs):
        return
    anchor = find_heading(doc, before_prefix)
    if not anchor:
        return
    idx = body_index_of_paragraph(anchor)
    insert_paragraph_before(doc, idx, heading_text, "Heading 1")


def append_supplement_sections(doc):
    doc.add_page_break()
    add_styled_paragraph(doc, "21. v1.5 补充规则与范围修订", "Heading 1")
    add_styled_paragraph(doc, "21.1 未来日期推荐范围", "Heading 2")
    doc.add_paragraph("第一版首页展示今天 + 未来 5 天天气。选中今天时，推荐入口正常可用；选中未来日期时，仅展示该日天气摘要和一句穿衣建议，不生成搭配推荐，推荐入口弱化并提示“第一版仅支持今日推荐”。")
    add_table_after(
        doc.paragraphs[-1],
        ["场景", "首页天气卡", "推荐入口状态", "推荐页行为"],
        [
            ("选中今天", "展示当日完整建议", "正常可点击", "基于今日天气生成推荐"),
            ("选中未来 5 天", "展示该日天气摘要", "弱化显示", "不生成推荐，提示今日才可获取推荐"),
        ],
    )
    add_styled_paragraph(doc, "21.2 当季风格标签定义", "Heading 2")
    doc.add_paragraph("第一版当季风格标签为系统预设，按当前月份判断季节，只展示当前季节 4-6 个标签。点击标签进入搭配推荐页，并将该风格作为筛选条件预填；若天气仍在获取中，标签点击同步禁用并提示“天气加载中，请稍候”。")
    add_table_after(
        doc.paragraphs[-1],
        ["季节", "预设风格标签"],
        [
            ("春", "清新通勤 / 法式复古 / 学院风 / 运动休闲 / 小香风"),
            ("夏", "度假风 / 简约清凉 / 美式休闲 / 甜酷风 / 森系"),
            ("秋", "复古文艺 / 知性通勤 / 街头潮流 / 奶咖色系 / 机车风"),
            ("冬", "高级感极简 / 甜系叠穿 / 街头大衣 / 韩系保暖 / 滑雪度假"),
        ],
    )
    add_styled_paragraph(doc, "21.3 推荐颜色协调逻辑", "Heading 2")
    doc.add_paragraph("第一版不做颜色硬过滤。原因是颜色协调规则复杂且 AI 颜色标签存在不确定性，过早过滤会减少推荐数量。第一版仅在推荐理由中说明颜色组合；若主色明显重复，可作为非强制排序调整。第二版再引入颜色搭配评分。")
    add_styled_paragraph(doc, "21.4 搭配预览降级说明", "Heading 2")
    doc.add_paragraph("第一版不实现固定人物底图叠图，改为 2×2 单品平铺卡片展示当套搭配。标题建议使用“搭配预览”，避免用户产生虚拟试穿预期。人物叠图/2D 虚拟模特能力移至第二版。")
    doc.add_paragraph("卡片布局：帽子（左上）/ 上衣（右上）/ 下装（左下）/ 鞋子（右下）。若某品类无推荐单品，对应卡片显示灰色占位块 + “暂无该品类”文字，不隐藏格子，不改变网格结构。")
    add_table_after(
        doc.paragraphs[-1],
        ["能力", "第一版", "第二版"],
        [
            ("搭配展示形式", "2×2 单品平铺卡片", "2D 虚拟模特换装"),
            ("AI 抠图", "不引入", "评估后引入"),
            ("合成预览图保存", "不生成", "评估后生成并存储"),
        ],
    )


def patch_recommendation_service(doc):
    heading = find_heading(doc, "20.11")
    if not heading:
        return
    start_idx = None
    end_idx = None
    for i, t in enumerate(doc.tables):
        joined = "\n".join(c.text for r in t.rows for c in r.cells)
        if "result.topId" in joined and "weatherDaily" in joined:
            start_idx = i
            break
    if start_idx is None:
        return
    # Replace the existing table with the multi-result structure.
    old_table = doc.tables[start_idx]
    prev_p = heading
    new_table = add_table_after(
        prev_p,
        ["字段", "类型", "说明"],
        [
            ("weatherDaily", "object", "当前选中日期的天气对象。"),
            ("wardrobeItems", "array", "当前用户衣橱列表。"),
            ("selectedDate", "string", "推荐日期。"),
            ("style", "string", "用户选择风格，可为空。"),
            ("occasion", "string", "用户选择场景，可为空。"),
            ("results", "array", "推荐结果列表，至少1套，最多3套。"),
            ("results[].topId", "string", "推荐上衣ID，可为空。"),
            ("results[].bottomId", "string", "推荐下装ID，可为空。"),
            ("results[].shoesId", "string", "推荐鞋子ID，可为空。"),
            ("results[].hatId", "string", "推荐帽子ID，可为空。"),
            ("results[].reason", "string", "推荐理由文案。"),
            ("results[].style", "string", "本套搭配的风格标签。"),
        ],
    )
    remove_table(old_table)
    # Update the rule paragraph.
    for p in doc.paragraphs:
        if p.text.startswith("规则：如果衣橱单品不足以生成完整搭配"):
            p.text = "规则：第一版推荐服务目标生成3套差异化方案；衣橱单品不足时至少返回1套，同时在 results[].reason 中说明缺少品类；衣橱为空时不执行推荐逻辑，直接返回空数组并附带引导文案。“换一组推荐”按钮触发时，重新执行推荐逻辑并打乱排列顺序，保证结果与上次不完全相同。"


def insert_after_paragraph(doc, anchor_prefix, additions):
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(anchor_prefix):
            anchor = p
            break
    if anchor is None:
        return
    idx = body_index_of_paragraph(anchor) + 1
    for text, style in additions:
        insert_paragraph_before(doc, idx, text, style)
        idx += 1


def append_table_after_heading(doc, heading_prefix, headers, rows, title=None):
    heading = find_heading(doc, heading_prefix)
    if not heading:
        return
    anchor = heading
    if title:
        idx = body_index_of_paragraph(heading) + 1
        anchor = insert_paragraph_before(doc, idx, title, "Heading 3")
    add_table_after(anchor, headers, rows)


def remove_paragraph(paragraph):
    el = paragraph._p
    el.getparent().remove(el)


def append_table_rows_if_missing(table, key_col, rows):
    existing = {row.cells[key_col].text.strip() for row in table.rows[1:]}
    for values in rows:
        if values[key_col] in existing:
            continue
        cells = table.add_row().cells
        for i, value in enumerate(values):
            set_cell_text(cells[i], value)
        existing.add(values[key_col])


def insert_table_row_after(table, anchor_key, key_col, values):
    existing = {row.cells[key_col].text.strip() for row in table.rows[1:]}
    if values[key_col] in existing:
        return
    new_row = table.add_row()
    for i, value in enumerate(values):
        set_cell_text(new_row.cells[i], value)
    for row in table.rows[1:]:
        if row.cells[key_col].text.strip() == anchor_key:
            row._tr.addnext(new_row._tr)
            return


def insert_paragraph_after_table(table, text, style=None):
    doc = table._parent
    p = doc.add_paragraph()
    if style:
        p.style = style
    r = p.add_run(text)
    set_font(r, 10.5, style and style.startswith("Heading"))
    table._tbl.addnext(p._p)
    return p


def find_table_by_header_and_key(doc, header, key_text):
    for table in doc.tables:
        if not table.rows:
            continue
        table_header = [cell.text.strip() for cell in table.rows[0].cells]
        if table_header != header:
            continue
        if any(key_text in cell.text for row in table.rows for cell in row.cells):
            return table
    return None


def replace_regex_everywhere(doc, pattern, repl):
    compiled = re.compile(pattern)
    for p in doc.paragraphs:
        new = compiled.sub(repl, p.text)
        if new != p.text:
            replace_in_runs(p, {p.text: new})
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    new = compiled.sub(repl, p.text)
                    if new != p.text:
                        replace_in_runs(p, {p.text: new})


def rebuild_delete_outfit_record_section(doc):
    start = find_heading(doc, "20.12b")
    end = find_heading(doc, "20.13")
    if not start or not end:
        return
    body = doc.element.body
    children = list(body)
    start_idx = children.index(start._p)
    end_idx = children.index(end._p)
    for child in children[start_idx:end_idx]:
        body.remove(child)

    idx = body_index_of_paragraph(end)
    heading = insert_paragraph_before(doc, idx, "20.12b deleteOutfitRecord 穿搭记录删除接口", "Heading 1")
    idx = body_index_of_paragraph(end)
    input_label = insert_paragraph_before(doc, idx, "入参", None)
    add_table_after(
        input_label,
        ["入参", "类型", "必填", "说明"],
        [("outfitId", "string", "是", "需要删除的穿搭记录ID")],
    )
    idx = body_index_of_paragraph(end)
    insert_paragraph_before(
        doc,
        idx,
        "删除规则：只能删除当前 openid 名下的记录；设置 isDeleted=true 和 deletedAt，不物理删除；outfitId 不存在或不属于当前用户时返回 RESOURCE_FORBIDDEN。",
        None,
    )
    idx = body_index_of_paragraph(end)
    return_label = insert_paragraph_before(doc, idx, "返回字段", None)
    add_table_after(
        return_label,
        ["返回字段", "类型", "说明"],
        [
            ("outfitId", "string", "已删除的穿搭记录ID"),
            ("deletedAt", "string", "软删除时间戳"),
        ],
    )


def patch_final_review_findings(doc):
    # D1: remove any leftover markdown comment marker before the architecture tree item.
    replace_everywhere(
        doc,
        {
            "#    ├─ saveOutfitRecord": "   ├─ saveOutfitRecord",
            "#   ├─ saveOutfitRecord": "   ├─ saveOutfitRecord",
            "# ├─ saveOutfitRecord": "├─ saveOutfitRecord",
        },
    )

    # N1: getWeather accepts lat/lng for automatic location.
    get_weather_params = find_table_by_header_and_key(
        doc,
        ["入参", "类型", "必填", "说明"],
        "cityName",
    )
    if get_weather_params:
        append_table_rows_if_missing(
            get_weather_params,
            0,
            [
                (
                    "lat",
                    "number",
                    "否",
                    "自动定位时传入纬度，与lng同时传入；有lat/lng时云函数优先做逆地理编码取城市名，忽略cityName。",
                ),
                ("lng", "number", "否", "自动定位时传入经度。"),
            ],
        )
        for row in get_weather_params.rows[1:]:
            if row.cells[0].text.strip() == "cityName":
                set_cell_text(row.cells[2], "否")
                set_cell_text(row.cells[3], "用户当前选择城市，如天津市。手动选城市时传入；自动定位传lat/lng时可为空。")

    # N2: align getWeather retry wording with section 7.5.
    replace_everywhere(
        doc,
        {
            "异常规则：天气获取失败时前端保持获取中状态，支持用户切换城市；推荐入口在成功前禁用或弱化。云函数可按固定间隔重试，但不得高频连续请求第三方接口。": "异常规则：天气获取失败时前端保持获取中状态，支持用户切换城市；推荐入口在成功前禁用或弱化。云函数重试策略与第7.5节一致：最多自动重试2次，首次等待2秒，再次等待5秒，三次均失败后返回 WEATHER_FAILED。",
        },
    )

    # N3: add outfit soft-delete fields and query filter.
    outfit_record_fields = find_table_by_header_and_key(
        doc,
        ["字段", "含义", "是否必填", "说明"],
        "outfitId",
    )
    if outfit_record_fields:
        append_table_rows_if_missing(
            outfit_record_fields,
            0,
            [
                ("isDeleted", "是否删除", "是", "软删除标记，默认 false"),
                ("deletedAt", "删除时间", "否", "软删除时写入"),
            ],
        )
    replace_everywhere(
        doc,
        {
            "主要字段：openid、date、topId、bottomId、shoesId、hatId、note、source、recommendationId、previewImageUrl、天气快照字段、createdAt、updatedAt": "主要字段：openid、date、topId、bottomId、shoesId、hatId、note、source、recommendationId、previewImageUrl、天气快照字段、isDeleted、deletedAt、createdAt、updatedAt",
            "查询规则：只返回当前 openid 名下记录；同一日期存在多条记录时按创建时间倒序返回。": "查询规则：只返回当前 openid 名下记录；默认只返回 isDeleted=false 的记录，已软删除的穿搭记录不展示在历史列表中；同一日期存在多条记录时按创建时间倒序返回。",
        },
    )

    # R1: add Tencent Map WebService API to external service plan and startup checklist.
    services = find_table_by_header_and_key(
        doc,
        ["服务", "用途", "当前状态", "处理原则"],
        "天气API",
    )
    if services:
        append_table_rows_if_missing(
            services,
            0,
            [
                (
                    "腾讯地图WebService API",
                    "逆地理编码，经纬度转城市名",
                    "待申请Key",
                    "通过getWeather云函数调用，Key存储在云函数环境变量中。",
                ),
            ],
        )
    insert_after_paragraph(
        doc,
        "申请天气API并记录服务商、Key、城市查询接口。",
        [("申请腾讯地图WebService API Key，确认逆地理编码接口权限。", None)],
    )

    # R2: remove resolved safety-rule item from pending questions.
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == "云数据库安全规则设计（建议仅允许用户读写自己的数据）。":
            remove_paragraph(paragraph)

    # R3: merge users fields into the main list and remove standalone note.
    replace_everywhere(
        doc,
        {
            "主要字段：openid、nickname、avatarUrl、defaultCity、themeMode、preferredStyles、preferredOccasions、createdAt、updatedAt": "主要字段：openid、nickname、avatarUrl、defaultCity、themeMode、preferredStyles、preferredOccasions、privacyConsentedAt、privacyPolicyVersion、dataDeleteRequestedAt、hasSkippedOnboarding、createdAt、updatedAt",
        },
    )
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == "users 集合补充字段：privacyConsentedAt、privacyPolicyVersion、dataDeleteRequestedAt、hasSkippedOnboarding。":
            remove_paragraph(paragraph)

    # R4: define deleteOutfitRecord return fields.
    if not find_table_by_header_and_key(doc, ["返回字段", "类型", "说明"], "deletedAt"):
        append_table_after_heading(
            doc,
            "20.12b deleteOutfitRecord",
            ["返回字段", "类型", "说明"],
            [
                ("outfitId", "string", "已删除的穿搭记录ID"),
                ("deletedAt", "string", "软删除时间戳"),
            ],
        )


def patch_second_review_findings(doc):
    # D1: regex fallback for any remaining markdown marker in the architecture tree.
    replace_regex_everywhere(
        doc,
        r"(?m)^(\s*)#\s*(├─\s*saveOutfitRecord\s+穿搭记录保存)",
        r"\1\2",
    )

    # 1: keep recommendationId flowing from recommendationService to outfit records.
    recommendation_output = find_table_by_header_and_key(
        doc,
        ["字段", "类型", "说明"],
        "results[].style",
    )
    if recommendation_output:
        insert_table_row_after(
            recommendation_output,
            "results[].style",
            0,
            (
                "results[].recommendationId",
                "string",
                "本套推荐的唯一ID，前端生成，格式建议 rec_{timestamp}_{index}，用于传入saveOutfitRecord做来源追溯。",
            ),
        )

    save_outfit_params = find_table_by_header_and_key(
        doc,
        ["入参", "类型", "必填", "说明"],
        "weatherSnapshot",
    )
    if save_outfit_params:
        for row in save_outfit_params.rows[1:]:
            key = row.cells[0].text.strip()
            if key == "weatherSnapshot":
                set_cell_text(
                    row.cells[3],
                    "记录当日天气快照，子字段与17.4节天气快照字段一致：weatherCity（城市）、weatherDate（天气日期）、weatherText（天气状态）、tempMin（最低温）、tempMax（最高温）、outfitSuggestion（穿衣建议文本），均为选填，前端从当前首页天气状态中取值传入。",
                )
        insert_table_row_after(
            save_outfit_params,
            "note",
            0,
            (
                "recommendationId",
                "string",
                "否",
                "source=recommendation时传入对应推荐ID；source=manual时可为空。",
            ),
        )

    # 4: define empty-city behavior for getWeather.
    get_weather_params = find_table_by_header_and_key(
        doc,
        ["入参", "类型", "必填", "说明"],
        "cityName",
    )
    if get_weather_params and "CITY_REQUIRED" not in "\n".join(p.text for p in doc.paragraphs):
        insert_paragraph_after_table(
            get_weather_params,
            "校验规则：cityName 与 lat/lng 至少提供一种；cityName、lat、lng 均为空时，云函数返回 CITY_REQUIRED，不发起天气请求。",
        )

    # 5: constrain outfit record dates in the page and cloud-function contract.
    replace_everywhere(
        doc,
        {
            "日期区：默认今日，支持修改为其他日期。": "日期区：默认今日，支持修改为今天及过去365天以内的日期；超出范围时禁用保存按钮并提示“只能记录近一年内的穿搭”。",
            "校验规则：保存穿搭记录时，topId/bottomId/shoesId/hatId至少有一个字段有值。同一日期允许保存多条记录，不覆盖已有记录。": "校验规则：保存穿搭记录时，topId/bottomId/shoesId/hatId至少有一个字段有值。同一日期允许保存多条记录，不覆盖已有记录。date 字段必须在今天及过去365天以内；超出范围时返回 DATE_OUT_OF_RANGE。",
        },
    )

    error_codes = find_table_by_header_and_key(
        doc,
        ["错误码", "含义", "前端处理"],
        "OK",
    )
    if error_codes:
        append_table_rows_if_missing(
            error_codes,
            0,
            [
                (
                    "CITY_REQUIRED",
                    "城市参数缺失，cityName与lat/lng均未传入",
                    "提示用户选择城市或开启位置权限，不发起天气请求。",
                ),
                (
                    "DATE_OUT_OF_RANGE",
                    "穿搭记录日期超出允许范围",
                    "提示用户选择近一年内的日期。",
                ),
            ],
        )

    # 3: rebuild the deleteOutfitRecord section in the same order as other APIs.
    rebuild_delete_outfit_record_section(doc)


def patch_security_and_integrity(doc):
    insert_after_paragraph(
        doc,
        "18.9.2 analyzeClothing",
        [
            ("安全校验：云函数收到 fileID 后，校验其路径是否符合 users/{当前openid}/clothing/ 前缀。路径不匹配时返回 RESOURCE_FORBIDDEN，不调用通义千问 VL API。", None),
            ("超时设置：云函数内对通义千问 VL API 的 HTTP 请求设置独立超时为 8 秒；8秒内未返回，云函数主动中止请求并返回 AI_TIMEOUT，不等待 VL API 响应。云函数整体超时在微信云开发控制台配置为 15 秒，为序列化、网络传输和图片读取预留余量。", None),
        ],
    )
    insert_after_paragraph(
        doc,
        "18.9.3 saveOutfitRecord",
        [
            ("安全校验：保存前，对入参中每个非空的 topId / bottomId / shoesId / hatId，在 clothing_items 集合中查询该ID是否存在且 openid 等于当前用户。任意一个ID不属于当前用户时，拒绝保存并返回 RESOURCE_FORBIDDEN，不写入记录。", None),
        ],
    )
    insert_after_paragraph(
        doc,
        "18.8 云存储方案",
        [
            ("孤立图片处理：图片上传成功后若 saveClothing 云函数调用失败，该图片成为孤立文件。第一版采用定期清理方案：后台每日扫描云存储中无对应 clothing_items 记录（或对应记录 isDeleted=true 且超过7天）的文件并删除。前端侧：saveClothing 失败时提示用户“保存失败，请重试”，不主动删除已上传图片，交由定期清理任务处理。", None),
        ],
    )
    insert_after_paragraph(
        doc,
        "18.11 天气接口接入建议",
        [
            ("逆地理编码方案：前端 wx.getLocation 获取经纬度后，将 lat/lng 传给 getWeather 云函数；云函数使用腾讯地图 WebService API 做逆地理编码，取 address_component.city 后再请求天气服务。腾讯地图 Key 存储在云函数环境变量中，不暴露前端。逆地理编码失败时，等同于位置权限拒绝，降级为手动选城市。", None),
            ("天气接口重试策略：云函数内自动重试最多2次；首次失败后等待2秒重试，再次失败后等待5秒重试，三次均失败后返回 WEATHER_FAILED。前端不做自动重试，展示“获取中”并等待用户主动切换城市重新触发。", None),
        ],
    )
    insert_after_paragraph(
        doc,
        "10.2 通用异常处理原则",
        [
            ("防重复提交：所有涉及数据写入的操作（新增衣物保存、穿搭记录保存），触发后立即将保存按钮置为禁用并展示 loading 状态，直至云函数返回结果后恢复。前端不依赖接口幂等，通过禁用按钮从源头阻止重复提交。", None),
        ],
    )
    insert_after_paragraph(
        doc,
        "5.5 超时与异常处理",
        [
            ("非服装图片处理：若 AI 返回的 category 字段与用户预选 categoryHint 不一致，或 confidence 低于 0.4 且 styleTags / subCategory 等关键字段均为空，前端在 AI 结果展示区顶部提示“AI未能识别该图片中的服装，建议手动填写标签”，不阻断保存流程，用户仍可手动填写后保存。", None),
            ("confidence 字段使用规则：confidence >= 0.4 时正常展示 AI 结果；confidence < 0.4 时，AI 结果展示区顶部展示提示“AI识别置信度较低，建议手动确认标签”。confidence 不影响保存流程，仅作为用户提示依据。", None),
        ],
    )
    append_table_after_heading(
        doc,
        "17.2",
        ["字段", "含义", "是否必填", "说明"],
        [
            ("privacyConsentedAt", "同意隐私政策时间", "是", "首次同意时写入，不可为空"),
            ("privacyPolicyVersion", "同意时的隐私政策版本号", "是", "隐私政策更新后重新触发确认时更新"),
            ("dataDeleteRequestedAt", "申请删除数据时间", "否", "用户注销或清除数据时写入"),
            ("hasSkippedOnboarding", "是否已跳过引导页", "否", "跳过冷启动引导后置为 true，默认 false"),
        ],
        "合规与冷启动字段",
    )
    insert_after_paragraph(
        doc,
        "18.7.1 users",
        [
            ("users 集合补充字段：privacyConsentedAt、privacyPolicyVersion、dataDeleteRequestedAt、hasSkippedOnboarding。", None),
            ("users 集合安全规则：{ \"read\": \"doc.openid == auth.openid\", \"write\": \"doc.openid == auth.openid\" }", None),
        ],
    )
    insert_after_paragraph(
        doc,
        "18.7.2 clothing_items",
        [
            ("clothing_items 集合安全规则：{ \"read\": \"doc.openid == auth.openid\", \"write\": \"doc.openid == auth.openid\" }", None),
        ],
    )
    insert_after_paragraph(
        doc,
        "18.7.3 outfit_records",
        [
            ("outfit_records 集合安全规则：{ \"read\": \"doc.openid == auth.openid\", \"write\": \"doc.openid == auth.openid\" }", None),
            ("上线前必须在微信云开发控制台逐集合配置以上规则，默认规则不可用于生产环境。", None),
        ],
    )
    # Add deleteOutfitRecord to cloud function list table.
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text == "名称":
            names = [row.cells[0].text for row in table.rows[1:]]
            if "saveOutfitRecord" in names and "deleteOutfitRecord" not in names:
                row = table.add_row().cells
                for i, value in enumerate(["deleteOutfitRecord", "云函数", "软删除指定穿搭记录", "是"]):
                    set_cell_text(row[i], value)
                break
    # getWardrobe filtering rule.
    insert_after_paragraph(
        doc,
        "分页规则：第一版可先按默认pageSize返回",
        [
            ("查询规则：默认只返回 isDeleted=false 的记录，不暴露已软删除衣物。dashboard.totalCount 和 categoryCounts 的统计口径与查询结果一致，仅计算 isDeleted=false 的有效记录。", None),
        ],
    )
    # deleteOutfitRecord interface after saveOutfitRecord.
    insert_after_paragraph(
        doc,
        "校验规则：保存穿搭记录时",
        [
            ("20.12b deleteOutfitRecord 穿搭记录删除接口", "Heading 1"),
        ],
    )
    append_table_after_heading(
        doc,
        "20.12b",
        ["入参", "类型", "必填", "说明"],
        [("outfitId", "string", "是", "需要删除的穿搭记录ID")],
    )
    insert_after_paragraph(
        doc,
        "20.12b deleteOutfitRecord",
        [
            ("删除规则：只能删除当前 openid 名下的记录；设置 isDeleted=true 和 deletedAt，不物理删除；outfitId 不存在或不属于当前用户时返回 RESOURCE_FORBIDDEN。", None),
        ],
    )
    # Add OUTFIT_NOT_FOUND error code.
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text == "错误码":
            codes = [row.cells[0].text for row in table.rows[1:]]
            if "OUTFIT_NOT_FOUND" not in codes:
                row = table.add_row().cells
                for i, value in enumerate(["OUTFIT_NOT_FOUND", "穿搭记录不存在或已删除。", "提示记录不存在，返回历史列表页。"]):
                    set_cell_text(row[i], value)
                break


def main():
    doc = Document(SRC)
    add_version_row(doc)

    replacements = {
        "v1.4  ·  审核稿": "v1.5  ·  审核稿",
        "本文件为 v1.4 审核稿": "本文件为 v1.5 审核稿",
        "主 PRD v1.4": "主 PRD v1.5",
        "人物预览": "搭配预览",
        "。；": "；",
        "人物搭配预览页": "搭配预览页",
        "人物搭配预览": "搭配预览",
        "人物搭配页": "搭配预览页",
        "人物搭配模块": "搭配预览模块",
        "人物搭配": "搭配预览",
        "固定人物底图 + 帽子/上衣/下装/鞋子叠加展示": "2×2 单品平铺卡片展示当套搭配组合（第一版；人物叠图移至第二版）",
        "第一版采用固定人物底图 + 单品叠加展示": "第一版采用 2×2 单品卡片平铺展示，不做人物叠图",
        "固定人物底图 + 单品图层叠加": "2×2 单品卡片平铺展示",
        "固定底图 + 帽子/上衣/下装/鞋子图层叠加，主链路必须项": "2×2 单品卡片平铺展示当套搭配组合，主链路必须项",
        "固定人物底图 + 单品图层叠加": "2×2 单品卡片平铺展示",
        "人物叠加配置": "搭配预览配置",
        "记录各单品叠加层级和位置": "记录各单品在预览卡片中的展示顺序",
        "人物叠层效果不理想": "人物叠图能力延期风险",
        "第一版先做规则化层级和位置，后续根据真实素材优化叠层模板。": "第一版不实现人物叠图，改为单品卡片预览；第二版评估图片质量和抠图方案后再引入。",
        "搭配预览第一版采用固定人物底图+单品图片叠层方案，不引入AI抠图。": "搭配预览第一版采用 2×2 单品平铺卡片展示，不引入 AI 抠图或人物叠图。",
        "在第一版固定人物底图基础上升级展示效果": "在第一版单品卡片预览基础上升级为正式换装展示效果",
        "用户可进入推荐页并查看搭配预览，至少生成 1 套完整搭配": "用户可进入推荐页并查看搭配预览；衣橱品类充足时至少生成 1 套完整搭配",
        "用户能看到固定人物底图上的搭配叠层预览，且可进入记录穿搭页。": "用户能看到 2×2 单品卡片形式的搭配预览，且可进入记录穿搭页。",
        "终止本次新增衣服流程": "终止本次新增衣物流程",
        "今日推荐入口卡（天气获取成功前禁用或弱化）。": "今日推荐入口卡（天气获取成功前禁用或弱化）。选中未来日期时，今日推荐入口卡切换为弱化状态，展示文案“第一版仅支持今日推荐”，点击不跳转推荐页；切换回今天后恢复正常可点击状态。",
        "人物展示主区域：固定底图叠加帽子/上衣/下装/鞋子单品图层。": "搭配预览主区域：2×2 单品卡片平铺，按帽子（左上）/ 上衣（右上）/ 下装（左下）/ 鞋子（右下）排列；若该品类无推荐单品，对应卡片显示灰色占位块 + “暂无该品类”文字，不隐藏格子。",
        "第一版共 3 个云函数，推荐逻辑不在云函数中，见第18.12节。": "第一版共 8 个云函数，推荐逻辑在前端本地服务中实现，不作为云函数，见第18.12节。云函数清单以第20.3节为准，本节仅展开三个核心函数的详细说明；其余五个（saveClothing、getWardrobe、updateClothing、deleteClothing、getOutfitRecords）的入参/出参定义见第20节对应小节。",
        "说明：第18部分已定义第一版关键云函数。本节为了支撑完整开发联调，对衣物CRUD和记录查询云函数进行补充。正式合并时需同步修订第18部分云函数清单，避免章节之间出现实现范围不一致。": "",
        "上传图片区：支持拍照或从相册选取单件衣物图。": "分类选择区（第一步）：进入新增衣物页后，优先展示分类选择“上衣 / 下装 / 鞋子 / 帽子”，用户选择后再展示上传图片区；所选分类作为 categoryHint 传入 AI 分析接口，辅助提升识别准确率。AI 返回分类与用户预选不一致时，以 AI 返回结果为建议结果，并允许用户手动修正。\n上传图片区（第二步）：支持拍照或从相册选取单件衣物图。",
        "流程顺序：上传图片 → AI 自动分析 → 用户确认 / 修改 → 补充选填字段 → 保存。": "流程顺序：1. 选择分类（必选） → 2. 上传单件衣物图片（拍照或相册，需上传媒体权限） → 3. AI 自动分析标签 → 4. 用户确认/修改分析结果 → 5. 用户补充基础字段与选填字段 → 6. 保存到衣橱。",
        "换一组推荐按钮：重新生成当前条件下的推荐结果。": "换一组推荐按钮：重新调用 recommendationService，生成新的最多3套方案，完全替换当前页全部推荐结果，等同于重置本轮推荐。",
        "操作按钮区：换一套（重新生成）、记录今日穿搭（带入当前搭配跳转记录页）、返回推荐列表。": "操作按钮区：换一套、记录今日穿搭（带入当前搭配跳转记录页）、返回推荐列表。换一套按钮在本轮已生成的多套推荐结果中切换至下一套（按顺序循环），不重新调用推荐服务；若本轮只有1套结果，换一套等同于换一组，重新调用推荐服务生成新结果。",
    }
    replace_everywhere(doc, replacements)

    # Specific table patches.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            if len(cells) >= 2 and cells[0] == "必填字段" and cells[1] == "分类、名称、图片":
                set_cell_text(row.cells[1], "分类、名称、图片（第一版必须上传图片才能保存；拒绝权限或取消 AI 分析后，均不可跳过图片直接录入）")
            if len(cells) >= 4 and cells[0] == "穿搭记录" and cells[1] == "历史回看、第二版统计":
                set_cell_text(row.cells[2], "用户主动创建，不主动操作则不产生此数据")
                set_cell_text(row.cells[3], "是（不记录即不产生）")
            if cells and cells[0] == "天气与首页":
                row.cells[1].text = row.cells[1].text + "；未来日期仅展示天气建议，不生成搭配推荐"
            if cells and cells[0] == "搭配推荐":
                row.cells[1].text = row.cells[1].text + "；第一版不做颜色硬过滤，仅在推荐理由中说明颜色组合"

    insert_section_15_16(doc)
    insert_12_9_12_10(doc)
    insert_top_heading_if_missing(doc, "17.1", "17. 数据结构设计")
    insert_18_4(doc)
    insert_top_heading_if_missing(doc, "18.1", "18. 技术架构建议")
    insert_top_heading_if_missing(doc, "19.1", "19. 实施计划")
    insert_top_heading_if_missing(doc, "20.1", "20. 接口与字段说明")
    append_supplement_sections(doc)
    patch_recommendation_service(doc)
    patch_security_and_integrity(doc)
    patch_final_review_findings(doc)
    patch_second_review_findings(doc)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            if len(cells) >= 4 and cells[0] == "穿搭记录" and cells[1] == "历史回看、第二版统计":
                set_cell_text(row.cells[2], "用户主动创建，不主动操作则不产生此数据")
                set_cell_text(row.cells[3], "是（不记录即不产生）")

    replace_everywhere(doc, {"。；": "；"})

    # Remove stray markdown comment marker if present.
    replace_everywhere(doc, {"#    ├─ saveOutfitRecord": "   ├─ saveOutfitRecord"})

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
