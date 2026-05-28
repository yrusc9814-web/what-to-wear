from copy import deepcopy
from pathlib import Path

from docx import Document


DOC_DIR = Path(r"D:\产品文档")
OFFICIAL = DOC_DIR / "穿搭日记微信小程序_PRD_v1.0.docx"
PART20 = DOC_DIR / "穿搭日记微信小程序_PRD_第20部分_接口与字段说明_v1.4_审核稿.docx"
FALLBACK = DOC_DIR / "穿搭日记微信小程序_PRD_v1.0_已加入第20部分.docx"


def element_text(el):
    return "".join(node.text or "" for node in el.iter() if node.tag.endswith("}t"))


def strip_existing_part20(doc):
    body = doc.element.body
    children = list(body)
    start_index = None
    for i, el in enumerate(children):
        if el.tag.endswith("}p") and element_text(el).strip().startswith("20.1 "):
            start_index = i
            break
    if start_index is None:
        return
    for el in children[start_index:]:
        if not el.tag.endswith("}sectPr"):
            body.remove(el)


def append_part20(base_doc, part_doc):
    strip_existing_part20(base_doc)
    base_doc.add_page_break()
    body = base_doc.element.body
    insert_before = body.sectPr
    include = False
    for el in part_doc.element.body:
        if el.tag.endswith("}sectPr"):
            continue
        if el.tag.endswith("}p") and element_text(el).strip().startswith("20.1 "):
            include = True
        if include:
            body.insert(body.index(insert_before), deepcopy(el))


def sync_part18_cloud_functions(doc):
    old_lines = {
        "   ├─ getWeather          天气接口中转": "   ├─ getWeather          天气接口中转",
        "   ├─ analyzeClothing     AI 标签分析中转（通义千问 VL）": "   ├─ analyzeClothing     AI 标签分析中转（通义千问 VL）",
        "   └─ saveOutfitRecord    穿搭记录保存": "   └─ saveOutfitRecord    穿搭记录保存",
    }
    new_block = [
        "   ├─ getWeather          天气接口中转",
        "   ├─ analyzeClothing     AI 标签分析中转（通义千问 VL）",
        "   ├─ saveClothing        衣物保存",
        "   ├─ getWardrobe         衣橱查询",
        "   ├─ updateClothing      衣物更新",
        "   ├─ deleteClothing      衣物软删除",
        "   ├─ saveOutfitRecord    穿搭记录保存",
        "   └─ getOutfitRecords    穿搭记录查询",
    ]

    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text == "   ├─ getWeather          天气接口中转":
            if i + 2 < len(paras) and paras[i + 1].text in old_lines and paras[i + 2].text in old_lines:
                for offset, text in enumerate(new_block):
                    if i + offset < len(paras):
                        paras[i + offset].text = text
                    else:
                        paras[i + offset - 1].insert_paragraph_before(text)
                # Remove leftover duplicated old lines if any remain immediately after the new block.
                extra_start = i + len(new_block)
                while extra_start < len(paras) and paras[extra_start].text in old_lines:
                    el = paras[extra_start]._element
                    el.getparent().remove(el)
                    paras = doc.paragraphs
                return True
    return False


def contains(path, marker):
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return marker in text


def main():
    if not OFFICIAL.exists():
        raise FileNotFoundError(OFFICIAL)
    if not PART20.exists():
        raise FileNotFoundError(PART20)

    base = Document(OFFICIAL)
    part20 = Document(PART20)
    synced = sync_part18_cloud_functions(base)
    append_part20(base, part20)

    try:
        base.save(OFFICIAL)
        target = OFFICIAL
    except PermissionError:
        base.save(FALLBACK)
        target = FALLBACK

    print(target)
    print("SYNCED_18", synced)
    print("HAS_20", contains(target, "20.1 本节定位"))
    print("HAS_UPDATE_CLOUD_FN", contains(target, "updateClothing      衣物更新"))


if __name__ == "__main__":
    main()
