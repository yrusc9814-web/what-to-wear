from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from pathlib import Path


OUTPUT_DIR = Path(r"D:\产品文档")
OUTPUT_FILE = OUTPUT_DIR / "穿搭日记微信小程序_PRD_v1.1_审核稿.docx"


def set_font(run, name="Microsoft YaHei", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_p(doc, text="", *, bold=False, size=11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    add_p(doc, text, bold=True, size=16 if level == 1 else 13)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)

    add_p(doc, "穿搭日记微信小程序", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "PRD v1.1 审核稿", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "说明：本文件为待确认审核版本，不替换现有正式 PRD", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "版本信息")
    add_bullets(doc, [
        "版本号：v1.1",
        "文档状态：审核稿",
        "说明：本版本在 v1.0 基础上完成结构重排、重复内容清理、版本边界明确、权限与异常补充。",
    ])

    add_heading(doc, "1. 产品背景与目标")
    add_p(doc, "1.1 产品背景", bold=True)
    add_p(doc, "随着年轻女性对日常形象管理、穿搭表达和生活方式记录的关注度不断提升，“今天穿什么”已经成为一个高频、持续、带有情绪影响的日常决策问题。")
    add_p(doc, "很多用户虽然拥有不少衣物，但在日常生活中仍然会遇到以下问题：")
    add_bullets(doc, [
        "出门前时间紧，但很难快速决定今天穿什么。",
        "衣橱利用率不高，很多单品长期闲置。",
        "面对天气变化、温差、下雨、换季时，难以快速判断怎么穿得舒适又得体。",
        "想记录穿搭，但现有方式零散，不利于长期沉淀。",
        "市面上很多穿搭内容更偏“看别人怎么穿”，而不是“我今天该怎么穿”。",
    ])
    add_p(doc, "微信小程序具备“低门槛、即开即用、适合高频轻决策”的特性，适合作为用户每天出门前快速打开使用的穿搭工具。因此，本产品希望以微信小程序为主要载体，打造一款围绕“数字衣橱、今日穿搭决策、每日穿搭记录”展开的轻量产品。")
    add_p(doc, "1.2 产品目标", bold=True)
    add_bullets(doc, [
        "解决即时问题：帮助用户在出门前更快决定今天穿什么。",
        "提升衣橱使用效率：通过数字化管理，让用户更清楚自己有什么、哪些衣物适合当前天气和场景。",
        "沉淀个人穿搭资产：通过长期记录每日穿搭，形成可回看、可复用、可优化的个人穿搭资料库。",
    ])
    add_p(doc, "1.3 产品价值主张", bold=True)
    add_p(doc, "这款产品对用户的核心价值不是“看更多穿搭内容”，而是：用你自己的衣橱，为你今天的穿搭做决定。")
    add_bullets(doc, [
        "省心：减少每天“今天穿什么”的决策负担。",
        "实用：结合天气、季节、场景和已有衣物，给出可落地的建议。",
        "沉淀：把每日穿搭积累成自己的数字衣橱和穿搭记录。",
    ])

    add_heading(doc, "2. 用户分析")
    add_p(doc, "2.1 核心目标用户", bold=True)
    add_p(doc, "18-30 岁女性用户，关注日常穿搭体验，愿意为“更省心、更好看、更适合自己”的穿搭决策投入少量时间，但不希望操作过重。")
    add_p(doc, "2.2 典型用户类型", bold=True)
    add_bullets(doc, [
        "学生用户：更关注风格感、记录感、约会和拍照表达，希望快速获得适合天气和心情的搭配建议。",
        "初入职场用户：更关注通勤效率和得体感，希望在早晨时间紧张的情况下快速得到不过度出错的穿搭方案。",
        "有记录习惯的穿搭爱好者：本身就有 OOTD 记录习惯，希望把穿搭做结构化沉淀，方便回看、复用和整理个人风格。",
    ])
    add_p(doc, "2.3 用户痛点", bold=True)
    add_bullets(doc, [
        "决策痛点：每天都可能陷入“今天穿什么”的高频决策焦虑。",
        "管理痛点：衣服很多，但并没有形成清晰的数字化整理。",
        "天气适配痛点：很难快速判断这种天气该怎么穿得舒服又好看。",
        "记录沉淀痛点：现有记录方式零散，不利于整理、搜索和复盘。",
        "个性化不足痛点：很多穿搭内容偏“看别人怎么穿”，缺少真正基于自己衣橱的建议。",
    ])

    add_heading(doc, "3. 核心使用场景")
    add_p(doc, "3.1 出门前快速决策场景", bold=True)
    add_bullets(doc, [
        "用户通常会在早上出门前、上课前、上班前、约会前打开小程序，希望在 1-3 分钟内快速解决“今天穿什么”的问题。",
        "产品需要提供：今日天气信息、一键查看今日推荐、根据当前季节给出风格引导、人物搭配预览、一键记录今日穿搭。",
    ])
    add_p(doc, "3.2 衣橱整理场景", bold=True)
    add_bullets(doc, [
        "用户会在周末、换季、买了新衣服之后集中把衣物录入系统，形成数字衣橱。",
        "产品需要提供：快速添加衣物、AI 自动补全标签、清晰的分类管理、扩展信息补充能力。",
    ])
    add_p(doc, "3.3 穿搭记录与回看场景", bold=True)
    add_bullets(doc, [
        "支持快速保存当日穿搭、查看历史记录和某日穿搭详情，形成可复用的穿搭资产库。",
    ])
    add_p(doc, "3.4 天气辅助决策场景", bold=True)
    add_bullets(doc, [
        "支持自动获取天气、手动选择城市、手动调整温度。",
        "根据天气输出穿衣建议，并让天气结果参与推荐逻辑。",
    ])

    add_heading(doc, "4. 产品定位")
    add_p(doc, "穿搭日记是一款面向年轻女性用户的微信小程序，围绕“数字衣橱、今日穿搭推荐、每日穿搭记录”三件事，帮助用户更轻松地做出日常穿搭决策，并逐步沉淀个人穿搭资产。")
    add_bullets(doc, ["轻量", "实用", "个性化", "可沉淀", "日常高频"])

    add_heading(doc, "5. MVP 范围与版本规划")
    add_p(doc, "5.1 第一版 MVP 核心目标", bold=True)
    add_p(doc, "第一阶段的目标，是先做出一个能在微信小程序中真实使用的工具型版本，验证用户是否愿意把它当成“每天决定穿什么”的工具来用。")
    add_p(doc, "5.2 第一版优先级原则", bold=True)
    add_bullets(doc, [
        "优先保证主链路闭环。",
        "优先做高频功能。",
        "优先做对推荐质量有直接帮助的功能。",
        "暂缓高复杂度但非第一阶段必要的能力。",
    ])
    add_p(doc, "5.3 第一版核心范围", bold=True)
    add_bullets(doc, [
        "首页天气与今日决策入口。",
        "数字衣橱管理。",
        "衣橱数据看板。",
        "AI 标签自动生成与手动修正。",
        "搭配推荐与人物搭配模块。",
        "同类衣物对比。",
        "每日穿搭记录与回看。",
        "浅色 / 深色 / 跟随系统主题切换。",
    ])
    add_p(doc, "5.4 第二版范围", bold=True)
    add_bullets(doc, [
        "完整分析报告页。",
        "身型与身材建议。",
        "风格雷达图。",
        "单品穿着次数统计。",
        "最久没穿分析。",
        "按月风格趋势图。",
        "包包 / 配饰扩展。",
        "更复杂的对比能力。",
        "2D 虚拟模特换装系统。",
    ])
    add_p(doc, "5.5 更后续规划", bold=True)
    add_bullets(doc, [
        "两件新上传待分析衣物之间对比。",
        "更复杂的用户画像与统计能力。",
        "更强的推荐逻辑和个性化学习。",
    ])

    add_heading(doc, "6. 第一版核心功能清单")
    add_bullets(doc, [
        "首页 / 今日决策入口。",
        "数字衣橱管理。",
        "衣橱数据看板。",
        "搭配推荐。",
        "人物搭配模块。",
        "同类衣物对比。",
        "每日穿搭记录。",
        "主题模式切换。",
    ])
    add_p(doc, "首页 / 今日决策入口包括：展示当前日期、当前城市天气、切换城市、手动调整温度、当天详细建议、未来 5 天日期入口、点击日期切换天气主卡、当季风格标签、今日推荐入口、快捷记录入口。")

    add_heading(doc, "7. AI 标签分析模块")
    add_p(doc, "7.1 功能目标", bold=True)
    add_p(doc, "用户上传单件衣物图后，系统自动识别并生成结构化衣物标签，帮助用户降低录入成本，提高衣橱数据完整度，并为后续推荐、筛选、对比提供基础数据支撑。")
    add_p(doc, "7.2 输入范围", bold=True)
    add_bullets(doc, [
        "第一版仅支持单件衣物图。",
        "第一版暂不支持商品截图、模特上身图和一张图中包含多件衣物。",
    ])
    add_p(doc, "7.3 输出方向", bold=True)
    add_bullets(doc, [
        "基础信息：大类、品类、适合性别。",
        "颜色信息：主色、次色、强调色、明度、色度。",
        "适用信息：季节适用性、适用场合、适合人群。",
        "设计信息：服装长度、合身度、材料、图案、细节特征。",
        "品类特征：领型、袖长、裤型、鞋型、帽型等。",
        "搭配辅助信息：风格标签、搭配关键词、建议搭配颜色、搭配难度、搭配提示。",
    ])
    add_p(doc, "7.4 编辑规则", bold=True)
    add_bullets(doc, [
        "AI 分析结果必须支持手动修改。",
        "AI 未识别出的字段允许为空。",
        "用户修改后的结果优先级高于 AI 输出。",
    ])
    add_p(doc, "7.5 异常处理", bold=True)
    add_bullets(doc, [
        "AI 分析失败时，提示失败并支持重新上传。",
        "AI 分析失败时，允许跳过 AI，直接进入手动录入流程。",
        "AI 返回字段较少时，不阻止用户保存。",
    ])

    add_heading(doc, "8. 衣橱录入与管理模块")
    add_p(doc, "8.1 功能目标", bold=True)
    add_p(doc, "帮助用户将现实衣物快速沉淀为数字衣橱，为推荐、对比、记录等功能提供稳定单品池。")
    add_p(doc, "8.2 字段设计", bold=True)
    add_bullets(doc, [
        "必填字段：分类、名称、图片。",
        "选填字段：备注、购入渠道、购入时间、购入价格、品牌名称、尺码、喜好程度、已穿次数。",
    ])
    add_p(doc, "8.3 主录入流程", bold=True)
    add_bullets(doc, [
        "上传单件衣物图片。",
        "AI 自动分析标签。",
        "用户确认 / 修改分析结果。",
        "用户补充基础与扩展字段。",
        "保存到衣橱。",
    ])
    add_p(doc, "8.4 管理能力", bold=True)
    add_bullets(doc, ["支持新增、编辑、删除、查看详情、分类浏览。"])
    add_p(doc, "8.5 异常处理", bold=True)
    add_bullets(doc, [
        "图片上传失败时，提示重新上传。",
        "删除衣物时需二次确认。",
        "AI 识别失败时可改为手动录入。",
    ])

    add_heading(doc, "9. 天气与首页决策模块")
    add_p(doc, "9.1 功能定位", bold=True)
    add_p(doc, "天气模块是首页最核心的信息模块之一，承担“告诉用户今天和未来几天怎么穿”的职责。")
    add_p(doc, "9.2 首页天气展示规则", bold=True)
    add_bullets(doc, [
        "第一版首页采用固定 6 天展示方案，即今天 + 后续 5 天。",
        "默认只展示当天日期、天气、温度区间和详细穿衣建议，默认选中“今天”。",
        "天气趋势条展示 6 天日期入口。",
        "用户点击某一天日期后，天气主卡整卡切换为该日对应信息。",
        "未选中日期只展示日期标签，不展开完整信息。",
    ])
    add_p(doc, "9.3 建议粒度", bold=True)
    add_bullets(doc, [
        "当天展示详细穿衣建议、厚薄建议、是否叠穿、是否建议带外套、帽子、防晒、防雨。",
        "未来 5 天展示日期、天气状态、温度区间和一句简短建议。",
    ])
    add_p(doc, "9.4 城市与温度能力", bold=True)
    add_bullets(doc, [
        "支持自动获取天气。",
        "支持手动切换城市。",
        "支持手动调整温度。",
        "所有推荐结果均以当前天气设定为依据。",
    ])
    add_p(doc, "9.5 异常处理", bold=True)
    add_bullets(doc, [
        "自动定位失败时，提示用户手动选择城市。",
        "天气接口失败时，允许用户继续使用手动城市和手动温度。",
        "即使自动天气不可用，用户仍可继续进入推荐流程。",
    ])

    add_heading(doc, "10. 搭配推荐与人物搭配模块")
    add_p(doc, "10.1 搭配推荐页", bold=True)
    add_bullets(doc, [
        "展示系统推荐的多个方案，帮助用户从推荐走向最终决策。",
        "显示精简版天气摘要，用于说明推荐依据。",
        "天气摘要仅展示当前城市、今日天气、温度区间和一句摘要提示，不重复首页完整天气模块。",
        "支持筛选条件、多套推荐结果、人物搭配预览入口、一键记录穿搭入口。",
    ])
    add_p(doc, "10.2 人物搭配页", bold=True)
    add_bullets(doc, [
        "第一版采用固定人物底图 + 单品叠加展示。",
        "展示帽子、上衣、下装、鞋子叠加效果。",
        "提供单品信息、搭配说明、换一套按钮和记录今日穿搭按钮。",
    ])

    add_heading(doc, "11. 服装对比模块")
    add_p(doc, "11.1 第一版范围", bold=True)
    add_bullets(doc, [
        "第一版优先支持衣橱内已上传衣物之间的同类对比。",
        "支持上衣 vs 上衣、下装 vs 下装、鞋子 vs 鞋子、帽子 vs 帽子。",
    ])
    add_p(doc, "11.2 交互规则", bold=True)
    add_bullets(doc, [
        "对比模式下最多只能选择两件。",
        "超过两件时不允许继续选择。",
        "第一版不支持跨分类对比。",
    ])
    add_p(doc, "11.3 对比维度", bold=True)
    add_bullets(doc, ["基础信息、颜色信息、适用信息、风格标签。"])
    add_p(doc, "11.4 后续版本归属", bold=True)
    add_bullets(doc, [
        "第二版：已上传衣物 vs 新上传待录入衣物对比。",
        "更后续版本：两件新上传待分析衣物之间对比。",
    ])

    add_heading(doc, "12. 主题模式模块")
    add_bullets(doc, [
        "第一版支持浅色模式、深色模式、跟随系统。",
        "用户可在设置页切换主题模式。",
        "主题状态需本地保存，并覆盖核心页面和组件。",
    ])

    add_heading(doc, "13. 页面结构与页面清单")
    add_p(doc, "13.1 底部导航结构", bold=True)
    add_bullets(doc, ["首页", "衣橱", "添加", "搭配", "我的"])
    add_p(doc, "13.2 “添加”入口定义", bold=True)
    add_p(doc, "底部中间的“添加”不是常规独立内容页，而是快捷操作入口。点击后弹出快捷菜单，提供“新增衣物”和“记录穿搭”两个选项。")
    add_p(doc, "13.3 页面清单", bold=True)
    add_bullets(doc, [
        "首页",
        "衣橱页",
        "衣物分类列表页",
        "新增衣物页",
        "衣物详情页",
        "衣物对比页",
        "搭配推荐页",
        "人物搭配预览页",
        "记录穿搭页",
        "穿搭历史页",
        "穿搭详情页",
        "我的页",
        "设置页",
    ])

    add_heading(doc, "14. 核心页面布局草案")
    add_p(doc, "14.1 首页", bold=True)
    add_bullets(doc, [
        "顶部基础区：标题、日期、问候语。",
        "今日天气主卡。",
        "六天天气趋势条。",
        "当季风格标签。",
        "今日推荐入口卡。",
        "快捷记录入口。",
    ])
    add_p(doc, "14.2 衣橱页", bold=True)
    add_bullets(doc, [
        "顶部标题区。",
        "数据看板区。",
        "分类导航区。",
        "最近新增 / 最近编辑区。",
    ])
    add_p(doc, "14.3 衣物分类列表页", bold=True)
    add_bullets(doc, [
        "顶部标题区：分类名称、数量、右上角“筛选”“对比”。",
        "筛选面板由顶部按钮触发，不常驻页面。",
        "第一版不做排序功能。",
        "第一版不做编辑功能。",
        "不使用底部操作条。",
        "顶部“筛选”位于“对比”左侧。",
    ])
    add_p(doc, "筛选面板包含：搜索、部位筛选、颜色筛选、风格筛选、季节筛选、关闭 / 重置 / 应用。")
    add_p(doc, "14.4 新增衣物页", bold=True)
    add_bullets(doc, ["上传图片区", "AI 分析状态区", "AI 标签结果区", "用户补充字段区", "保存操作区"])
    add_p(doc, "14.5 搭配推荐页", bold=True)
    add_bullets(doc, ["精简版天气摘要区", "筛选条件区", "推荐结果区", "换一组推荐按钮"])
    add_p(doc, "14.6 人物搭配预览页", bold=True)
    add_bullets(doc, ["搭配标题区", "人物展示主区域", "单品信息区", "搭配说明区", "操作按钮区"])
    add_p(doc, "14.7 记录穿搭页", bold=True)
    add_bullets(doc, ["日期区", "搭配结果区", "备注区", "保存区"])
    add_p(doc, "补充说明：首页“快捷记录入口”点击后进入记录穿搭页。从首页直接进入时，用户手动选择单品并完成记录；从推荐页或人物搭配页进入时，自动带入当前推荐搭配结果。")
    add_p(doc, "14.8 我的页与设置页", bold=True)
    add_bullets(doc, [
        "我的页：穿搭历史入口、设置入口、主题模式入口。",
        "设置页：浅色模式、深色模式、跟随系统。",
    ])

    add_heading(doc, "15. 权限与异常处理")
    add_p(doc, "15.1 权限申请时机", bold=True)
    add_bullets(doc, [
        "相册 / 拍照权限：在用户点击上传衣物图片时申请。",
        "位置权限：在用户首次使用自动天气能力时申请。",
    ])
    add_p(doc, "15.2 权限拒绝后的降级方案", bold=True)
    add_bullets(doc, [
        "相册 / 拍照权限拒绝后：提示用户开启权限，支持后续重新触发授权，若无法拍照则继续保留相册上传入口或走手动录入策略。",
        "位置权限拒绝后：提示用户手动选择城市，允许用户手动调整温度，不阻断推荐流程。",
    ])
    add_p(doc, "15.3 通用异常处理原则", bold=True)
    add_bullets(doc, [
        "所有异常都需有用户可理解的提示文案。",
        "所有关键流程都应提供重试或降级路径。",
        "异常不应让主链路完全中断。",
    ])

    add_heading(doc, "16. 版本变更记录")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["版本", "日期", "变更内容", "备注"]
    for i, text in enumerate(headers):
        run = hdr[i].paragraphs[0].add_run(text)
        set_font(run, bold=True)
    rows = [
        ("v1.0", "当前已整理版本", "微信小程序 PRD 初稿", "已落文档"),
        ("v1.1", "本次审核稿", "重构章节、去重、补权限与异常、明确版本边界、明确添加入口交互", "待确认"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            run = cells[i].paragraphs[0].add_run(text)
            set_font(run)

    add_heading(doc, "17. 当前结论")
    add_bullets(doc, [
        "以天气驱动的今日穿搭决策为主线。",
        "以数字衣橱和 AI 标签分析为基础。",
        "以人物搭配展示和每日穿搭记录作为核心体验闭环。",
        "以衣物对比和主题模式作为增强体验能力。",
        "第二版再补完整分析页、统计分析和更强虚拟换装能力。",
    ])

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(str(OUTPUT_FILE))


if __name__ == "__main__":
    build()
