from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from pathlib import Path


OUTPUT_DIR = Path(r"D:\产品文档")
OUTPUT_FILE = OUTPUT_DIR / "穿搭日记_产品文档_v0.1.docx"


def set_font(run, name="Microsoft YaHei", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph(document, text="", style=None, bold=False, size=11):
    p = document.add_paragraph(style=style)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True)
    return p


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run)


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("穿搭日记 ✦ OOTD\n产品文档 v0.1")
    set_font(run, size=18, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于当前 index.html 原型整理")
    set_font(run, size=10)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("代码目录：D:\\chuanda-project    产品文档目录：D:\\产品文档")
    set_font(run, size=10)

    add_heading(doc, "1. 产品概述", level=1)
    add_paragraph(doc, "产品名：穿搭日记 ✦ OOTD")
    add_paragraph(
        doc,
        "产品定位：面向女性用户的轻量级日常穿搭助手，帮助用户管理衣橱、记录每日穿搭、结合天气获取搭配建议，并沉淀个人穿搭数据。"
    )
    add_paragraph(
        doc,
        "产品阶段判断：当前更接近高保真静态原型 / 前端可交互 Demo，还不是可正式上线的小程序版本。"
    )
    add_paragraph(doc, "核心价值：")
    add_bullets(
        doc,
        [
            "降低“今天穿什么”的决策成本。",
            "让用户把散乱衣物整理成可组合的数字衣橱。",
            "用“记录”和“回顾”增强使用习惯。",
            "用天气和场景风格做轻量辅助推荐。",
        ],
    )

    add_heading(doc, "2. 目标用户", level=1)
    add_paragraph(doc, "核心用户：18-30 岁、日常有穿搭表达需求的女生。")
    add_paragraph(doc, "典型人群：")
    add_bullets(
        doc,
        [
            "上学/上班前纠结穿什么的用户。",
            "喜欢记录 daily look、OOTD 的用户。",
            "衣服不少，但不会高频组合搭配的用户。",
            "想把衣橱数字化、方便回看和管理的用户。",
        ],
    )

    add_heading(doc, "3. 当前产品结构", level=1)
    add_paragraph(doc, "当前页面结构分为 3 个主 Tab：")
    add_bullets(
        doc,
        [
            "首页：今日日期、记录今日穿搭入口、穿搭日记列表。",
            "衣橱页：上衣 / 下装 / 鞋子分类管理、试穿台、AI 抠图入库。",
            "AI 推荐页：天气卡片、手动天气兜底、季节筛选、风格筛选、推荐结果卡片。",
        ],
    )
    add_paragraph(doc, "另有 4 个弹层：")
    add_bullets(
        doc,
        [
            "记录穿搭弹层。",
            "添加衣物弹层。",
            "AI 抠图弹层。",
            "数据管理弹层。",
        ],
    )

    add_heading(doc, "4. 当前已实现功能", level=1)
    add_paragraph(doc, "4.1 首页 / 穿搭日记", bold=True)
    add_bullets(
        doc,
        [
            "展示当天日期。",
            "新增、编辑、删除穿搭记录。",
            "每条记录包含日期、上衣、下装、鞋子、备注、缩略图预览。",
            "空状态引导用户创建第一套穿搭。",
        ],
    )
    add_paragraph(doc, "4.2 衣橱管理", bold=True)
    add_bullets(
        doc,
        [
            "按上衣 / 下装 / 鞋子分类查看。",
            "手动添加衣物。",
            "上传图片并自动压缩。",
            "删除衣物。",
            "衣物支持名称与备注。",
        ],
    )
    add_paragraph(doc, "4.3 试穿台", bold=True)
    add_bullets(
        doc,
        [
            "点击衣物卡片即可试穿。",
            "分别挂载到上衣/下装/鞋子槽位。",
            "用 SVG 小人做实时穿搭预览。",
            "可清空单件或全部清空。",
            "可从试穿结果一键进入“记录穿搭”。",
        ],
    )
    add_paragraph(doc, "4.4 AI 抠图", bold=True)
    add_bullets(
        doc,
        [
            "上传全身照或服装图。",
            "本地压缩后进入分析。",
            "通过前景背景分离做简化去底。",
            "将整张图分成上衣/下装/鞋子三个区域。",
            "用户可拖动滑块手动调整切分位置。",
            "每个区域可命名并保存到衣橱。",
        ],
    )
    add_paragraph(doc, "说明：这里的“AI 抠图”本质上是前端 Canvas 图像处理 + 简单规则分区，不是真正接入模型服务。")
    add_paragraph(doc, "4.5 天气模块", bold=True)
    add_bullets(
        doc,
        [
            "优先浏览器定位获取天气。",
            "若失败则尝试 IP 定位。",
            "若仍失败则降级到手动输入。",
            "展示温度、天气、城市、状态徽章。",
            "根据天气给出一句轻量穿搭提示。",
            "手动天气模式可被导出备份并导入恢复。",
        ],
    )
    add_paragraph(doc, "4.6 AI 推荐", bold=True)
    add_bullets(
        doc,
        [
            "支持季节筛选：全季 / 春 / 夏 / 秋 / 冬。",
            "支持风格筛选：不限 / 休闲 / 通勤 / 甜妹 / 酷飒。",
            "基于现有衣橱生成 3 套推荐。",
            "推荐卡包含搭配标签、衣物列表、小人预览、推荐理由、天气适配标记。",
            "可将推荐结果一键带入“记录穿搭”。",
        ],
    )
    add_paragraph(doc, "说明：当前“AI 推荐”并不是调用大模型，而是基于现有衣物池 + 随机/规则排序 + 天气权重的推荐机制。")
    add_paragraph(doc, "4.7 数据管理", bold=True)
    add_bullets(
        doc,
        [
            "展示数据概览。",
            "导出 JSON 备份。",
            "导入 JSON 恢复。",
            "清空全部本地数据。",
        ],
    )

    add_heading(doc, "5. 当前交互主流程", level=1)
    add_bullets(
        doc,
        [
            "用户先往衣橱里添加衣物。",
            "可直接点衣物做试穿预览。",
            "满意后记录成当天穿搭。",
            "或进入 AI 推荐页，结合天气生成建议。",
            "选中推荐后回填到记录页保存。",
            "长期积累后，在首页回看穿搭日记。",
        ],
    )

    add_heading(doc, "6. 当前数据设计", level=1)
    add_paragraph(doc, "当前数据完全存储在浏览器本地 localStorage 中。")
    add_paragraph(doc, "主要数据对象：")
    add_bullets(
        doc,
        [
            "clothes：id / type / name / note / img / createdAt。",
            "outfits：id / date / topId / bottomId / shoesId / note / createdAt / updatedAt。",
            "weatherPrefs：仅在手动天气模式下备份 temp / cond / city。",
        ],
    )

    add_heading(doc, "7. 当前产品亮点", level=1)
    add_bullets(
        doc,
        [
            "体验完整度比普通 Demo 高，主链路比较顺。",
            "视觉风格统一，明显偏女性化审美。",
            "“衣橱 - 试穿 - 记录”这一段很顺，是产品最有感觉的部分。",
            "天气降级设计比较细，不容易直接报错卡死。",
            "支持导出导入，已经考虑用户资产。",
        ],
    )

    add_heading(doc, "8. 当前问题与风险", level=1)
    add_bullets(
        doc,
        [
            "“AI”概念容易被用户误解，当前能力更接近规则推荐与图像分区。",
            "推荐逻辑依赖衣橱结构，但衣物标签太少，缺少颜色、厚薄、场合、季节、版型等字段。",
            "generateAI() 调用了 getWeatherFitScore() 和 buildWeatherReason()，但当前文件中未发现对应函数定义，AI 推荐功能可能存在运行错误风险。",
            "数据只存本地浏览器，不适合正式产品长期使用。",
            "缺少场景维度，如通勤、约会、出游、下雨天等。",
            "缺少长期激励机制，如打卡、复盘、偏好画像等。",
        ],
    )

    add_heading(doc, "9. 版本结论", level=1)
    add_paragraph(
        doc,
        "从现状看，这个产品已经具备较完整的原型雏形，最适合作为后续小程序产品化的基础版本。当前最有价值的资产不是“AI”本身，而是已经跑通的衣橱管理、试穿预览、穿搭记录和天气辅助推荐链路。"
    )
    add_paragraph(
        doc,
        "备注：后续产品更新后，建议持续同步维护本文件，作为当前版本产品定义的基准文档。"
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(str(OUTPUT_FILE))


if __name__ == "__main__":
    build_document()
